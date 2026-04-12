"""
FastAPI backend for GCSE Worksheet QA Studio
Handles document processing, agent validation, and export operations
"""

import os
import re
import json
import base64
import secrets
import asyncio
import hashlib
from io import BytesIO
from typing import Optional
from pathlib import Path

from fastapi import FastAPI, UploadFile, File, Form, Request
from fastapi.responses import StreamingResponse, FileResponse, Response
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_TAB_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openai import OpenAI, AsyncOpenAI

# Import agent prompts
import sys
from pathlib import Path

# Add parent directory to path to import agents.py
sys.path.insert(0, str(Path(__file__).parent.parent))

from agents import (
    FORMATTING_AGENT_PROMPT,
    AGENT_1_PROMPT,
    AGENT_2_PROMPT,
    AGENT_3_PROMPT,
    AGENT_4_PROMPT,
    AGENT_5_PROMPT,
)

# ============================================================================
# CONFIGURATION
# ============================================================================

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
async_client = AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Simple in-memory export cache — keyed by SHA-256 of text content.
# Avoids re-running the formatting agent on every download of the same text.
_export_cache: dict = {}

ANSWER_LINE = "_____________________________________________________________________________"
ANSWER_UNDERSCORES = {0: 79, 1: 76, 2: 73}
LABEL_CM = [0.0, 0.63, 1.27]
TEXT_CM = [0.7, 1.27, 1.90]

# ============================================================================
# FASTAPI APP
# ============================================================================

app = FastAPI(title="GCSE Worksheet QA Studio")

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Mount static files for frontend
frontend_path = Path(__file__).parent.parent / "frontend"
if frontend_path.exists():
    app.mount("/static", StaticFiles(directory=str(frontend_path)), name="static")


# ── HTTP Basic Auth middleware ────────────────────────────────────────────────
# Only active when AUTH_PASSWORD environment variable is set.
# Local dev with no AUTH_PASSWORD set → no auth required.
@app.middleware("http")
async def basic_auth_middleware(request: Request, call_next):
    auth_password = os.getenv("AUTH_PASSWORD", "")
    if not auth_password:
        # No password configured — allow all (local dev)
        return await call_next(request)

    auth_user = os.getenv("AUTH_USER", "examqa")
    auth_header = request.headers.get("Authorization", "")

    if auth_header.startswith("Basic "):
        try:
            decoded = base64.b64decode(auth_header[6:]).decode("utf-8")
            username, _, password = decoded.partition(":")
            user_ok = secrets.compare_digest(username.strip(), auth_user)
            pass_ok = secrets.compare_digest(password.strip(), auth_password)
            if user_ok and pass_ok:
                return await call_next(request)
        except Exception:
            pass

    return Response(
        status_code=401,
        headers={"WWW-Authenticate": 'Basic realm="examqa"'},
        content="Unauthorized — please enter your credentials.",
    )


# ============================================================================
# HELPER FUNCTIONS
# ============================================================================


def _set_run_font(run, bold=False, size_pt=11):
    """Set font properties for a run in a DOCX document."""
    run.font.name = "Arial"
    run.font.size = Pt(size_pt)
    run.bold = bold
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rPr.insert(0, rFonts)


def extract_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file."""
    doc = Document(BytesIO(file_bytes))
    return "\n".join([p.text for p in doc.paragraphs])


def clean_text(text: str) -> str:
    """Remove markdown-style formatting characters."""
    return re.sub(r"[#*]+", "", text)


def add_answer_lines(text: str) -> str:
    """Add underscore answer lines after questions with marks."""
    lines = text.split("\n")
    output = []
    for line in lines:
        output.append(line)
        match = re.search(r"\((\d+)\)", line)
        if match:
            marks = int(match.group(1))
            for _ in range(min(marks, 4)):
                output.append(ANSWER_LINE)
    return "\n".join(output)


def extract_total(text: str) -> Optional[int]:
    """Extract total marks from worksheet text."""
    match = re.search(r"Total for paper\s*=\s*(\d+)", text)
    return int(match.group(1)) if match else None


def fractional_marks_present(text: str) -> bool:
    """Check if worksheet contains fractional marks."""
    return bool(re.search(r"\(\d+\.\d+\)", text))


def keyword_overlap(text1: str, text2: str) -> float:
    """Calculate keyword overlap percentage between two texts."""
    words1 = set(re.findall(r"\b[a-zA-Z]{5,}\b", text1.lower()))
    words2 = set(re.findall(r"\b[a-zA-Z]{5,}\b", text2.lower()))
    if not words1:
        return 0
    return round((len(words1 & words2) / len(words1)) * 100, 1)


def extract_question_numbers(text: str) -> list:
    """Extract main question numbers from worksheet."""
    nums = set()
    for m in re.finditer(r"^\s*(\d+)\s*(?=[.(A-Za-z])", text, re.MULTILINE):
        n = m.group(1)
        try:
            v = int(n)
        except ValueError:
            continue
        if v <= 50:
            nums.add(n)
    return sorted(nums, key=lambda x: int(x))


def strip_answer_lines(text: str) -> str:
    """Remove underscore answer lines from text."""
    lines = text.split("\n")
    return "\n".join([ln for ln in lines if ANSWER_LINE.strip() not in ln.strip()])


def detect_question_structure(text: str) -> list:
    """Detect the hierarchical structure of questions and sub-parts."""
    ROMAN_RE = re.compile(r"^\s*\((i{1,4}|iv|vi{0,3}|ix|xi{0,3}|x{1,3})\)\s", re.IGNORECASE)
    PART_RE = re.compile(r"^\s*\(([a-z])\)\s")
    structure = {}
    current_q = None
    current_part = None
    for line in text.split("\n"):
        if "Total for question" in line:
            continue
        m_main = re.match(r"^\s*(\d+)\s*(?=[.(A-Za-z])", line)
        if m_main:
            v = int(m_main.group(1))
            if v <= 50:
                current_q = m_main.group(1)
                current_part = None
                structure.setdefault(current_q, {"parts": {}})
            continue
        if current_q is None:
            continue
        m_roman = ROMAN_RE.match(line)
        if m_roman and current_part is not None:
            roman = m_roman.group(1).lower()
            structure[current_q]["parts"].setdefault(current_part, set())
            structure[current_q]["parts"][current_part].add(roman)
            continue
        m_part = PART_RE.match(line)
        if m_part:
            letter = m_part.group(1)
            structure[current_q]["parts"].setdefault(letter, set())
            current_part = letter
    result = []
    for qnum, info in structure.items():
        parts_list = []
        for letter in sorted(info["parts"].keys()):
            roman_set = sorted(info["parts"][letter])
            entry = {"letter": letter}
            if roman_set:
                entry["roman_subparts"] = roman_set
            parts_list.append(entry)
        result.append({"question_number": qnum, "parts": parts_list})
    return result


def read_spec_text(
    spec_txt_bytes: Optional[bytes] = None,
    spec_docx_bytes: Optional[bytes] = None,
    pasted_spec_text: Optional[str] = None,
) -> str:
    """Combine specification text from multiple sources."""
    parts = []
    if spec_txt_bytes:
        try:
            parts.append(spec_txt_bytes.decode("utf-8"))
        except Exception:
            pass
    if spec_docx_bytes:
        try:
            parts.append(extract_docx(spec_docx_bytes))
        except Exception:
            pass
    if pasted_spec_text:
        parts.append(pasted_spec_text)
    return "\n\n".join(p.strip() for p in parts if p and p.strip())


# ============================================================================
# AI FUNCTIONS
# ============================================================================


async def improve_worksheet(text: str, spec_text: str = "") -> str:
    """Improve worksheet quality and formatting using AI."""
    spec_block = ""
    if spec_text.strip():
        spec_block = f"""
SPECIFICATION CONTENT — SCOPE CONSTRAINT (CRITICAL):
{spec_text.strip()}

You MUST stay strictly within this specification when improving or rewriting questions.
- Do NOT introduce any scientific fact, equation, definition, or value that is not in the specification above.
- When rewriting poor-quality questions, draw ONLY from the specification content.
- When adding application scenarios, base them only on content present in the specification.
- If a question's topic is not in the specification, rewrite it around a topic that IS.

"""
    prompt = f"""
You are improving a GCSE science worksheet to match a professional exam-standard format.
The worksheet may be for any GCSE science subject: Biology, Chemistry, Physics, or combined science.
Apply the same standards regardless of subject. Follow these rules EXACTLY:
{spec_block}
CONTENT RULES:
1. Make every question clear and unambiguous. Remove AI-sounding or strange wording.
2. Questions should replicate real GCSE exam questions in style and difficulty.
3. Ensure a MIX of question types: 1-mark recall, 2-3 mark describe, 3-4 mark explain, calculation questions.
4. Include at least one application-style question with a named scenario (e.g. "A student, Sarah, connects a circuit...").
5. Do NOT repeat questions or topics.
6. Ensure every multi-mark question has appropriate cognitive demand.
7. Numbers in question stems should be written as WORDS (e.g. "two" not "2"), EXCEPT for: physical values (e.g. "15 m/s"), equations, or units.

COMMAND WORD RULES (CRITICAL):
8. Every question instruction — whether a main question or a sub-part — MUST start with a GCSE command word.
   Accepted command words: Explain, State, Describe, Calculate, Determine, Show that, Identify, Give, Name,
   Suggest, Compare, Evaluate, Predict, Draw, Plot, Label, Complete, Define, Outline, Justify, Use, Write.
9. NEVER phrase a question instruction as an interrogative sentence starting with "What", "Which", "Why",
   "How", "When" or "Where". Convert these to command-word form:
     - "What is the unit of force?" → "State the unit of force."
     - "Which type of wave is used in optical fibres?" → "Identify the type of wave used in optical fibres."
     - "Why does the wave slow down?" → "Explain why the wave slows down."
     - "How does a concave mirror improve the beam?" → "Explain how a concave mirror improves the beam."
10. LINE SEPARATION — CRITICAL: Every standalone statement or sentence that makes sense on its own MUST
    be on its OWN line. This applies to EVERY part of the worksheet without exception.
    Specifically:
    - Context / scenario sentences MUST be on a line of their own, NEVER merged onto the same line as the
      command-word instruction that follows them.
    - Each command-word instruction MUST be on its own line.
    - NEVER write context and instruction as a single run-on line.

    CORRECT (context and command each on their own line):
        (b)  Radio waves are used for broadcasting.
             Explain why radio waves can be received over large distances.   (2)

    INCORRECT (merged onto one line — NEVER do this):
        (b)  Radio waves are used for broadcasting. Explain why radio waves can be received over large distances.   (2)

    This rule applies to EVERY question throughout the entire worksheet, not just some of them.
    Apply it consistently to questions 1 through the last question.

FORMATTING RULES:
11. Remove ALL topic headers (e.g. "Work and Energy Transfers", "Forces", "Section A").
12. Remove ALL formatting symbols: *, #, bullet points, dashes used as headers.
13. In any question context/stem text, replace " = " and " - " used as label separators with ": ".
    Example: "Work Done = Force x Distance" in a context line -> "Work done: force x distance"
14. Question numbering must be consistent: 1, 2, 3 ... (a), (b), (c) ... (i), (ii), (iii).
    - Main question numbers should NOT have a dot (use "1" not "1.")
15. Every main question number (1, 2, 3...) MUST include an introductory sentence of its own BEFORE any
    sub-parts. Never leave a main question number bare with no text on its line.
    Example of CORRECT format:
        3  A student investigates refraction of light at a glass block.
           (a) State what happens to the speed of light as it enters the glass.   (1)
           (b) Explain why the ray changes direction at the boundary.   (2)
    Example of INCORRECT format (bare question number — NEVER do this):
        3
           (a) State what happens to the speed of light...
16. Do NOT add answer lines - these are handled separately.
17. Keep mark allocations exactly as shown, e.g. (2).
18. Ensure there is NO space between sub-parts (a), (b), (c) of the SAME question.
19. There SHOULD be a blank line between separate main questions (1, 2, 3...).
20. REWRITE PERMISSION: If a question is clearly of poor quality — contains vague wording, is not GCSE-appropriate, uses interrogative phrasing throughout, or cannot be fixed with minor edits — you MUST rewrite it substantially to produce a high-quality GCSE question. Do not preserve poor-quality content for the sake of preservation. Clarity, GCSE realism, and exam authenticity always take priority.
21. If a question has sub-parts (a)(i), (a)(ii), the letter (a) alone should NOT be on its own line
    if it only introduces roman-numeral sub-parts. Use the format:
    (a) (i) question text here   (1)
        (ii) question text here  (2)

IMAGE PLACEHOLDERS:
22. If a question references a diagram, figure, graph, or image, do NOT describe it inline.
    Insert a placeholder on its own line in exactly this format:
        [Image: <description of what the image/diagram should show>]
    The description must be specific enough that a teacher can source or draw the correct image.
    - CORRECT: [Image: Diagram showing a root hair cell with a large vacuole and thin cell wall]
    - WRONG: [Image: diagram here]
    Placeholders are rendered as visual boxes in the preview and replaced manually when printing.

OUTPUT:
Return the improved worksheet only. No commentary or explanations.
"""
    response = await async_client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": prompt},
            {"role": "user", "content": text},
        ],
        temperature=0,
    )
    return add_answer_lines(clean_text(response.choices[0].message.content))



async def generate_markscheme(text: str, spec_text: str = "", mismatch_info: Optional[str] = None) -> str:
    """Generate a mark scheme from the worksheet text."""
    mismatch_block = ""
    if mismatch_info:
        mismatch_block = f"""
CRITICAL - SPECIFIC ISSUES TO FIX IN THIS REGENERATION:
{mismatch_info}
You MUST resolve every issue listed above. Do not reproduce these errors.
"""

    spec_block = ""
    if spec_text.strip():
        spec_block = f"""
SPECIFICATION CONTENT (use this to guide mark scheme answers):
{spec_text.strip()}

CRITICAL SPEC RULE: Every marking point must be traceable to the specification content above.
Use the exact terminology, definitions, equations, and values from the specification.
Do NOT introduce scientific facts, definitions, or values that are not present in the specification.
"""

    prompt = f"""
You are generating a fully explicit GCSE-style mark scheme from a science worksheet.
The worksheet may cover Biology, Chemistry, Physics, or combined science — apply the same mark scheme standards for all.
{mismatch_block}{spec_block}

CONTENT RULES:
1. Write a marking entry for EVERY question and EVERY sub-part in the worksheet.
2. NEVER skip any sub-question.
3. NEVER use vague placeholders like "Working step (1)" or "Answer (1)".
4. For CALCULATION questions:
   - Show the actual equation, numerical substitution, and final answer with units.
   - Do NOT award a mark for giving the equation alone (guideline rule).
   - Award marks for correct substitution (1) and correct answer with units (1).
   - Example: "a = (v - u) / t = (0 - 20) / 5 = -4 m/s2 (1)(1)"
5. For NON-CALCULATION questions:
   - Give clear, specific marking points.
   - Allow reasonable alternatives: "OR" / "Accept..."
   - If 3+ possible answers exist, use: "Any [one/two/three] from:" followed by bullet points.
   - Use "OR" when only two alternatives exist.
6. Each mark MUST be a whole number - use (1) only. Never (2) for a single point.
7. Only the FIRST letter of each marking sentence should be capitalised.
8. Sentences longer than 3-4 words MUST end with a full stop before the (1).
9. Any useful side note should be prefixed with [NOTE]:
   e.g. "[NOTE]: Accept velocity instead of speed"

FORMATTING RULES:
10. Bold question numbers: "1", "2" etc. (just the number).
11. Sub-part labels in brackets: (a), (b), (c), (i), (ii).
12. NO space between marking points WITHIN the same question part.
13. A SMALL space (one blank line) between DIFFERENT sub-parts (a)->(b)->(c).
14. A SMALL space between the last mark of one question and the Total line.
15. Include a "(Total for question X is Y marks)" line after each main question.
    CRITICAL: Use "is" NOT "=" — e.g. "(Total for question 1 is 6 marks)" NOT "(Total for question 1 = 6 marks)"
16. At the very END of the mark scheme, add on its own line:
    "Total marks for question paper: Z"
    where Z is the sum of all question totals.

Question mapping and numbering:
- Use EXACTLY the question numbers and sub-parts from the DETECTED QUESTION STRUCTURE below.
- Do NOT invent new question numbers or sub-parts.
- Do NOT omit any question or sub-part.
- Treat numbers inside sentences (e.g. "2.0 m/s", "5 kg") as data NOT question numbers.
- Only start a new question number at the beginning of a line.
"""

    response = await async_client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": prompt},
            {
                "role": "user",
                "content": (
                    "WORKSHEET TEXT:\n"
                    + text
                    + "\n\nDETECTED QUESTION STRUCTURE (do NOT change these IDs):\n"
                    + json.dumps(detect_question_structure(text), ensure_ascii=False)
                ),
            },
        ],
        temperature=0,
    )
    return clean_text(response.choices[0].message.content)


async def improve_markscheme(worksheet_text: str, existing_ms: str, spec_text: str = "") -> str:
    """Improve and validate an uploaded mark scheme against the worksheet."""
    spec_block = ""
    if spec_text.strip():
        spec_block = f"""
SPECIFICATION CONTENT (use this to guide mark scheme answers):
{spec_text.strip()}

CRITICAL SPEC RULE: Every marking point must be traceable to the specification content above.
Use the exact terminology, definitions, equations, and values from the specification.
Do NOT introduce scientific facts or values that are not present in the specification.

"""
    prompt = f"""You are reviewing and improving an uploaded GCSE science mark scheme against its worksheet.
The worksheet may cover Biology, Chemistry, Physics, or combined science — apply the same standards for all subjects.
{spec_block}
CONTENT RULES:
1. Keep all correct marking points — do NOT discard good content.
2. Fix any vague marking points (e.g. "correct answer (1)", "working step (1)") — replace with the actual answer, equation, or value.
3. For CALCULATION questions ensure:
   - Show the actual equation, numerical substitution, and final answer with units.
   - Do NOT award a mark purely for writing an equation — marks are for substitution and correct answer.
   - Example: "a = (v - u) / t = (0 - 20) / 5 = -4 m/s² (1)(1)"
4. For NON-CALCULATION questions give clear, specific marking points.
   - When 3+ possible answers exist, use "Any [one/two/three] from:" followed by bullet points.
   - Use "OR" only when exactly two alternatives exist on the same line.
5. Every individual mark must use (1). Never use (2) or (3) for a single marking point.
6. Every marking point must be a SHORT, concise phrase — not a full paragraph.

STRUCTURE RULES:
7. Ensure every question and sub-part in the worksheet has a corresponding mark scheme entry. Add any missing ones.
8. Do NOT invent question numbers or sub-parts that don't exist in the worksheet.
9. Only the FIRST letter of each marking sentence should be capitalised.
10. Sentences longer than 3-4 words must end with a full stop before the (1).
11. Any accept/note guidance should be on its own line: [NOTE]: accept X instead of Y

FORMATTING RULES:
12. Bold question numbers: "1", "2" etc. (just the number).
13. Sub-part labels in brackets: (a), (b), (c), (i), (ii).
14. No space between marking points WITHIN the same question part.
15. One blank line between DIFFERENT sub-parts (a) → (b) → (c).
16. One blank line between the last mark of a question and the Total line.
17. Include "(Total for question X is Y marks)" after each main question — use "is" NOT "=".
18. At the very end, on its own line: "Total marks for question paper: Z"
19. Remove ALL markdown symbols (#, *, etc.).

Return the improved mark scheme only. No commentary or explanations.
"""
    response = await async_client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": prompt},
            {
                "role": "user",
                "content": (
                    "WORKSHEET:\n" + worksheet_text
                    + "\n\nUPLOADED MARK SCHEME TO IMPROVE:\n" + existing_ms
                ),
            },
        ],
        temperature=0,
    )
    return clean_text(response.choices[0].message.content)


async def run_agent(prompt: str, content: str, model: str = "gpt-4o-mini") -> str:
    """Run an agent. Retries up to 2 times on transient errors with exponential back-off."""
    last_err: Exception = RuntimeError("No attempts made")
    for attempt in range(3):
        try:
            response = await async_client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": prompt},
                    {"role": "user", "content": content},
                ],
                temperature=0,
            )
            return response.choices[0].message.content
        except Exception as e:
            last_err = e
            if attempt < 2:
                await asyncio.sleep(2 ** attempt)  # 1 s, then 2 s
    raise last_err



def parse_revised_output(text: str) -> tuple:
    """Parse Agent 5 output to extract revised worksheet and mark scheme."""
    ws_marker = "--- REVISED WORKSHEET ---"
    ms_marker = "--- REVISED MARK SCHEME ---"
    ws_idx = text.find(ws_marker)
    ms_idx = text.find(ms_marker)
    if ws_idx != -1 and ms_idx != -1:
        return text[ws_idx + len(ws_marker) : ms_idx].strip(), text[ms_idx + len(ms_marker) :].strip()
    ws_match = re.search(r"-{2,}\s*REVISED\s+WORKSHEET\s*-{2,}", text, re.IGNORECASE)
    ms_match = re.search(r"-{2,}\s*REVISED\s+MARK\s+SCHEME\s*-{2,}", text, re.IGNORECASE)
    if ws_match and ms_match:
        return text[ws_match.end() : ms_match.start()].strip(), text[ms_match.end() :].strip()
    return None, None


async def run_formatting_agent(worksheet_text: str) -> dict:
    """Run formatting agent to get JSON structure for document building."""
    cleaned = strip_answer_lines(clean_text(worksheet_text))
    response = await async_client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": FORMATTING_AGENT_PROMPT},
            {"role": "user", "content": cleaned},
        ],
        temperature=0,
    )
    raw = response.choices[0].message.content.strip()
    raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.IGNORECASE)
    raw = re.sub(r"\s*```$", "", raw)
    try:
        return json.loads(raw)
    except json.JSONDecodeError as exc:
        raise ValueError(
            f"FormattingAgent returned invalid JSON ({exc}). "
            f"First 500 chars:\n{raw[:500]}"
        ) from exc


# ============================================================================
# DOCUMENT BUILDING FUNCTIONS
# ============================================================================


def build_formatted_docx(spec: dict) -> BytesIO:
    """
    Build a fully formatted A4 Word document matching AQA GCSE exam paper style.

    Layout:
      • Question text on its own line with label (no inline marks)
      • Answer lines as underscore-character paragraphs below the question
      • Marks (n) on a separate RIGHT-ALIGNED BOLD paragraph after answer lines
      • (Total for question X is Y marks) — right-aligned bold
    """
    document = Document()

    # Page setup (A4, 1.91cm side margins)
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(1.91)
    section.right_margin = Cm(1.91)

    # Normal style
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    lines = spec.get("lines", [])
    paper_total = spec.get("paper_total_marks")
    last_q = None
    prev_indent = None
    _COMB_RE_D = re.compile(r"^\(([a-z])\) \(([ivxlcdm]+)\)$", re.IGNORECASE)
    _seen_roman_d: dict = {}
    _orphan_qnum_d = None

    def _para(space_before_pt=0, space_after_pt=0, left_cm=0.0, hanging_cm=0.0, align=WD_ALIGN_PARAGRAPH.LEFT):
        p = document.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(space_before_pt)
        pf.space_after = Pt(space_after_pt)
        pf.alignment = align
        if left_cm:
            pf.left_indent = Cm(left_cm)
        if hanging_cm:
            pf.first_line_indent = Cm(-hanging_cm)
        return p

    _CMD_SPLIT_D = re.compile(
        r"(?<=[.?!])\s+(Calculate|Explain|State|Describe|Determine|Show|Find|"
        r"Compare|Evaluate|Suggest|Identify|Give|Write|Draw|Plot|Predict|Justify|"
        r"Define|Outline|Use|Work\s+out|Complete|Name|Tick|Circle|Underline|Label)\b",
        re.IGNORECASE,
    )

    for line in lines:
        qnum = line.get("question_number")
        indent_level = int(line.get("indent_level", 0))
        part_label = line.get("part_label") or ""
        subpart_label = line.get("subpart_label") or ""
        question_text = line.get("question_text") or ""
        if question_text.strip().lower() == "none":
            question_text = ""
        marks = line.get("marks")
        is_total = bool(line.get("is_total_for_question"))

        # Total for question
        if is_total:
            p = _para(space_before_pt=4, space_after_pt=12, align=WD_ALIGN_PARAGRAPH.RIGHT)
            r = p.add_run(f"(Total for question {qnum} is {marks} marks)")
            _set_run_font(r, bold=True)
            last_q = qnum
            prev_indent = None
            _orphan_qnum_d = None
            continue

        # Level-0 items with no text — save number to merge onto next line
        if indent_level == 0 and not question_text.strip():
            _orphan_qnum_d = qnum
            last_q = qnum
            continue

        # Roman numeral de-duplication
        _eff_indent = indent_level
        _eff_part = part_label
        if indent_level == 1:
            _mcd = _COMB_RE_D.match(part_label)
            if _mcd:
                _parent_d = _mcd.group(1)
                _roman_d = _mcd.group(2)
                _seen_roman_d.setdefault(qnum, set())
                if _parent_d in _seen_roman_d[qnum]:
                    _eff_indent = 2
                    _eff_part = f"({_roman_d})"
                else:
                    _seen_roman_d[qnum].add(_parent_d)

        # Merge orphan question number
        _ans_lvl_d = _eff_indent
        if _orphan_qnum_d is not None:
            _eff_part = (f"{_orphan_qnum_d}  {_eff_part}" if _eff_part else str(_orphan_qnum_d))
            _eff_indent = 0
            _orphan_qnum_d = None

        # Spacing before this paragraph
        if _eff_indent == 0:
            sp = 22 if (last_q is not None and qnum != last_q) else 0
        elif _eff_indent == 1:
            sp = 10 if (prev_indent is not None and prev_indent != 0) else 6
        else:  # level 2
            sp = 8

        # Label (plain text — no bold per AQA exam style)
        if _eff_indent == 0:
            label = _eff_part or (str(qnum) if qnum else "")
        elif _eff_indent == 1 and _eff_part:
            label = _eff_part
        elif _eff_indent >= 2 and (subpart_label or _eff_part):
            label = subpart_label or _eff_part
        else:
            label = ""
        label_bold = False

        # Question text paragraph
        lvl = min(_eff_indent, 2)
        _child_lvl = min(_ans_lvl_d, 2)
        if lvl == 0 and _child_lvl > 0:
            # Was a merged orphan — wrap at child TEXT_CM
            _left_cm = TEXT_CM[_child_lvl]
            _hang_cm = TEXT_CM[_child_lvl]
        else:
            _left_cm = TEXT_CM[lvl]
            _hang_cm = TEXT_CM[lvl] - LABEL_CM[lvl]

        # Auto-split context sentence from command word, then handle \n splits
        question_text = _CMD_SPLIT_D.sub(lambda m: "\n" + m.group(0).lstrip(), question_text)
        _qt_parts = [s for s in question_text.split("\n") if s.strip()]

        _txt_cm = TEXT_CM[min(_eff_indent, 2)]

        if len(_qt_parts) > 1:
            # Context line(s) first (no label), then label + command on the last line.
            # This puts "(b)" next to "Evaluate this statement." not next to the context sentence.
            for _ci, _ctx in enumerate(_qt_parts[:-1]):
                _ctx_sp = sp if _ci == 0 else 0
                _cp = _para(space_before_pt=_ctx_sp, space_after_pt=0, left_cm=_txt_cm)
                _set_run_font(_cp.add_run(_ctx), bold=False)
            # Label + instruction paragraph
            qp = _para(space_before_pt=0, space_after_pt=0, left_cm=_left_cm, hanging_cm=_hang_cm)
            if label:
                _set_run_font(qp.add_run(label + "  "), bold=label_bold)
            _set_run_font(qp.add_run(_qt_parts[-1]), bold=False)
        else:
            # Single part — label + text together as before
            qp = _para(space_before_pt=sp, space_after_pt=0, left_cm=_left_cm, hanging_cm=_hang_cm)
            if label:
                _set_run_font(qp.add_run(label + "  "), bold=label_bold)
            if _qt_parts:
                _set_run_font(qp.add_run(_qt_parts[0]), bold=False)

        # Answer lines (underscore text)
        if marks and marks > 0:
            m = int(marks)
            num_ans = min(m, 6)
            _alvl = min(_ans_lvl_d, 2)
            underscores = "_" * ANSWER_UNDERSCORES.get(_alvl, 62)

            for i in range(num_ans):
                ap = _para(space_before_pt=0, space_after_pt=0, left_cm=LABEL_CM[_alvl])
                r = ap.add_run(underscores)
                _set_run_font(r, bold=False)

            # Marks on their own right-aligned line
            mp = _para(space_before_pt=0, space_after_pt=2, left_cm=LABEL_CM[_alvl], align=WD_ALIGN_PARAGRAPH.RIGHT)
            r_m = mp.add_run(f"({marks})")
            _set_run_font(r_m, bold=True)

        last_q = qnum
        prev_indent = _eff_indent

    # Paper total
    if paper_total:
        p = _para(space_before_pt=14, align=WD_ALIGN_PARAGRAPH.RIGHT)
        r = p.add_run(f"Total marks for question paper: {paper_total}")
        _set_run_font(r, bold=True)
        r.underline = True

    bio = BytesIO()
    document.save(bio)
    bio.seek(0)
    return bio


def build_markscheme_docx(markscheme_text: str) -> BytesIO:
    """
    Build a mark scheme DOCX matching the Edexcel format:
    - Each marking point ends with bold (1)
    - 'Any X from:' lines followed by bullet points
    - (Total for question X is Y marks) — bold, 'is' not '='
    - Total marks line — bold + underlined
    """
    document = Document()
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(1.91)
    section.right_margin = Cm(1.91)

    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    MAIN_Q_RE = re.compile(r"^\s*(\d+)\s")
    PART_RE = re.compile(r"^\s*\(([a-z])\)")
    ROMAN_RE = re.compile(r"^\s*\((i{1,4}|iv|vi{0,3}|ix|xi{0,3}|x{1,3})\)", re.IGNORECASE)
    BULLET_RE = re.compile(r"^\s*[•\-\*]\s+")
    TOTAL_Q_RE = re.compile(r"\(Total for question", re.IGNORECASE)
    TOTAL_PAPER_RE = re.compile(r"Total marks for question paper", re.IGNORECASE)
    TOTAL_PAPER_ALT = re.compile(r"Total for paper\s*[=:]", re.IGNORECASE)

    def add_inline_bold_marks(p, text, base_bold=False):
        """Split text on (1) and emit alternating normal / bold runs."""
        parts = re.split(r"(\(1\))", text)
        for part in parts:
            if part == "(1)":
                r = p.add_run("(1)")
                _set_run_font(r, bold=True)
            elif part:
                r = p.add_run(part)
                _set_run_font(r, bold=base_bold)

    prev_line_type = None

    for raw_line in markscheme_text.split("\n"):
        line = raw_line.strip()
        if not line:
            continue

        # Total for paper
        if TOTAL_PAPER_RE.search(line) or TOTAL_PAPER_ALT.search(line):
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(line)
            _set_run_font(r, bold=True)
            r.underline = True
            prev_line_type = "total"
            continue

        # Total for question
        if TOTAL_Q_RE.search(line):
            # Normalise "=" -> "is"
            line = re.sub(
                r"(Total for question\s+\w+)\s*=\s*(\d+)",
                r"\1 is \2",
                line,
                flags=re.IGNORECASE,
            )
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(12)
            r = p.add_run(line)
            _set_run_font(r, bold=True)
            prev_line_type = "total"
            continue

        # Bullet point
        if BULLET_RE.match(line):
            text_after_bullet = BULLET_RE.sub("", line).strip()
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Cm(1.8)
            p.paragraph_format.first_line_indent = Cm(-0.4)
            r = p.add_run("\u2022 ")
            _set_run_font(r, bold=False)
            add_inline_bold_marks(p, text_after_bullet)
            prev_line_type = "bullet"
            continue

        # Identify line type
        is_main = MAIN_Q_RE.match(line)
        is_part = PART_RE.match(line) or ROMAN_RE.match(line)

        if is_main:
            sp_before = 14 if prev_line_type in ("total", "part", "bullet", "other") else 0
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(sp_before)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Cm(0)
            m = is_main
            r = p.add_run(m.group(1) + " ")
            _set_run_font(r, bold=True)
            add_inline_bold_marks(p, line[m.end() :])
            prev_line_type = "main"

        elif is_part:
            sp_before = 4 if prev_line_type in ("part", "bullet", "other") else 0
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(sp_before)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Cm(0.8)
            add_inline_bold_marks(p, line)
            prev_line_type = "part"

        else:
            # Continuation / any-from line / notes
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Cm(1.5)
            add_inline_bold_marks(p, line)
            prev_line_type = "other"

    bio = BytesIO()
    document.save(bio)
    bio.seek(0)
    return bio


# ============================================================================
# API ENDPOINTS
# ============================================================================


def _truthy(v) -> bool:
    """Return True if v represents a meaningful non-zero/non-false value."""
    if v is None: return False
    if isinstance(v, bool): return v
    if isinstance(v, (int, float)): return v > 0
    s = str(v).strip().lower()
    return s not in ("", "0", "false", "none", "no", "n/a", "na", "null")


def _count_agent_issues(r1: str, r2: str, r3: str, r4: str) -> dict:
    """
    Parse all four agent JSON reports and return counts plus a flat list of
    human-readable insight strings for the Pipeline Insights panel.
    Each insight has: {text: str, level: "fix"|"warn"|"info"}
    """
    counts = {"command_word": 0, "mark_scheme": 0, "cognitive": 0, "coverage": 0, "total": 0}
    insights = []   # list of {text, level}

    # ── Agent 1 — command words & interrogative phrasing ─────────────────────
    try:
        d = json.loads(r1)
        interrogative, under, over, missing = 0, 0, 0, 0
        for q in d.get("question_analysis", []):
            st = q.get("alignment_status", "ok")
            if q.get("is_interrogative") or st == "interrogative":
                interrogative += 1
            elif st == "under_rewarded":
                under += 1
            elif st == "over_rewarded":
                over += 1
            elif st == "missing":
                missing += 1
        counts["command_word"] = interrogative + under + over + missing
        if interrogative:
            insights.append({"text": f"{interrogative} question{'s' if interrogative>1 else ''} used interrogative phrasing (What/Why/How) instead of a command word — reworded", "level": "fix"})
        if missing:
            insights.append({"text": f"{missing} question instruction{'s' if missing>1 else ''} had no command word — added", "level": "fix"})
        if under:
            insights.append({"text": f"{under} question{'s' if under>1 else ''} under-rewarded for cognitive demand (e.g. Explain worth only 1 mark)", "level": "warn"})
        if over:
            insights.append({"text": f"{over} question{'s' if over>1 else ''} over-rewarded for cognitive demand (e.g. State worth 4+ marks)", "level": "warn"})
    except Exception:
        pass

    # ── Agent 2 — mark scheme structural violations ───────────────────────────
    try:
        d = json.loads(r2)
        viol = d.get("summary", {}).get("rule_violations", {})

        gran  = viol.get("granularity_errors", "")
        fmt   = viol.get("total_line_format_errors", "")
        paper = viol.get("paper_total_missing", False)
        vague = viol.get("vague_marking_points", "")
        caps  = viol.get("capitalisation_errors", "")
        anyx  = viol.get("any_x_from_errors", "")

        # count mark-mismatches from per-question analysis
        mismatch = sum(1 for q in d.get("question_analysis", []) if q.get("structure_alignment") == "mismatch")

        if _truthy(gran):
            insights.append({"text": f"Mark granularity errors found — individual marks were grouped as (2) or (3) instead of separate (1)s", "level": "fix"})
            counts["mark_scheme"] += 1
        if _truthy(fmt):
            insights.append({"text": "Question totals used '=' instead of 'is' — corrected to exam board format", "level": "fix"})
            counts["mark_scheme"] += 1
        if _truthy(paper):
            insights.append({"text": "Paper total mark line was missing from mark scheme — added", "level": "fix"})
            counts["mark_scheme"] += 1
        if _truthy(vague):
            insights.append({"text": "Vague marking points found (e.g. 'correct answer (1)') — replaced with explicit expected answers", "level": "fix"})
            counts["mark_scheme"] += 1
        if _truthy(caps):
            insights.append({"text": "Capitalisation inconsistencies in mark scheme — only first letter of each point should be capitalised", "level": "fix"})
            counts["mark_scheme"] += 1
        if _truthy(anyx):
            insights.append({"text": "'Any X from:' format was missing or malformed — standardised", "level": "fix"})
            counts["mark_scheme"] += 1
        if mismatch:
            insights.append({"text": f"{mismatch} question{'s' if mismatch>1 else ''} had mark totals that didn't match the number of (1) mark labels", "level": "fix"})
            counts["mark_scheme"] += mismatch

        calc_q = d.get("summary", {}).get("calculation_marking_quality", "")
        if calc_q and any(w in calc_q.lower() for w in ("poor", "issue", "error", "wrong", "incorrect")):
            insights.append({"text": "Calculation marking quality flagged — marks were awarded for writing equations rather than substitution/answer", "level": "fix"})
            counts["mark_scheme"] += 1
    except Exception:
        pass

    # ── Agent 3 — cognitive balance & GCSE authenticity ──────────────────────
    try:
        d = json.loads(r3)
        bf = d.get("balance_flags", {})
        auth = d.get("authenticity_issues", {})
        rating = d.get("overall_quality_rating", "")
        dist = d.get("mark_distribution", {})

        if bf.get("too_much_recall"):
            pct = dist.get("recall_pct", "")
            insights.append({"text": f"Paper was recall-heavy{' (' + str(pct) + ')' if pct else ''} — cognitive balance adjusted toward explain/application questions", "level": "fix"})
            counts["cognitive"] += 1
        if bf.get("no_calculations"):
            insights.append({"text": "No calculation questions detected — paper lacked procedural skills assessment", "level": "warn"})
            counts["cognitive"] += 1
        if bf.get("no_causal_or_extended"):
            insights.append({"text": "No causal or extended-response questions — paper was missing higher-order thinking", "level": "warn"})
            counts["cognitive"] += 1
        if bf.get("heavy_skew_detected"):
            skew = bf.get("skew_description", "")
            insights.append({"text": f"Cognitive skew detected{': ' + skew if skew else ''}", "level": "warn"})
            counts["cognitive"] += 1

        if not auth.get("difficulty_gradient_ok", True):
            insights.append({"text": "Difficulty gradient broken — hard questions appeared before easier ones; reordered", "level": "fix"})
            counts["cognitive"] += 1

        ai_lang = auth.get("ai_sounding_language_found", "")
        if _truthy(ai_lang):
            insights.append({"text": f"AI-generated phrasing detected (e.g. 'delve into', 'fascinating') — replaced with natural exam language", "level": "fix"})
            counts["cognitive"] += 1

        mismatches = auth.get("mark_cognitive_mismatches", "")
        if _truthy(mismatches):
            insights.append({"text": "Mark–cognitive mismatches found (e.g. Recall question worth 4 marks, Explain worth 1)", "level": "fix"})
            counts["cognitive"] += 1

        scenario_q = auth.get("scenario_quality_issues", "")
        if _truthy(scenario_q):
            insights.append({"text": "Named-scenario questions had unrealistic or physically incoherent setups — corrected", "level": "fix"})
            counts["cognitive"] += 1

        if rating in ("poor", "fair"):
            insights.append({"text": f"Overall cognitive quality rated '{rating}' before revision — significant structural improvements applied", "level": "info"})
    except Exception:
        pass

    # ── Agent 4 — topic coverage / spec scope ────────────────────────────────
    try:
        d = json.loads(r4)
        out_of_scope = [q for q in d.get("per_question", []) if q.get("scope_status", "in_scope") != "in_scope"]
        counts["coverage"] = len(out_of_scope)
        if out_of_scope:
            qids = ", ".join(q.get("question_id", "?") for q in out_of_scope[:4])
            extra = f" (and {len(out_of_scope)-4} more)" if len(out_of_scope) > 4 else ""
            insights.append({"text": f"{len(out_of_scope)} question{'s' if len(out_of_scope)>1 else ''} outside the specification removed or rewritten — Q{qids}{extra}", "level": "fix"})

        gaps = d.get("coverage_gaps", "")
        if _truthy(gaps) and len(str(gaps)) > 5:
            insights.append({"text": f"Spec topics not covered by the worksheet: {gaps}", "level": "info"})

        overrep = d.get("overrepresented_topics", "")
        if _truthy(overrep) and len(str(overrep)) > 5:
            insights.append({"text": f"Overrepresented topics: {overrep}", "level": "info"})
    except Exception:
        pass

    counts["total"] = counts["command_word"] + counts["mark_scheme"] + counts["cognitive"] + counts["coverage"]
    return {**counts, "_insights": insights}


async def process_stream_generator(
    worksheet_bytes: bytes,
    markscheme_bytes: Optional[bytes],
    spec_text: str,
):
    """Async generator for streaming processing steps.

    Pipeline:
      0 — Read & parse files
      1 — Enhance worksheet  (gpt-4o-mini)
      2 — Build/improve mark scheme  (gpt-4o-mini)
      3 — Agents 1-4 run IN PARALLEL  (gpt-4o-mini × 4 concurrent)
      4 — Agent 5 final revision  (gpt-4o — higher quality for the critical rewrite)

    Any unhandled exception yields an error event so the frontend can surface a
    message rather than leaving the progress spinner running indefinitely.
    """
    try:
        step = 0

        # ── Step 0: Read files ────────────────────────────────────────────────────
        yield f'data: {json.dumps({"step": step, "label": "Reading files", "detail": "Parsing documents"})}\n\n'
        worksheet_text = extract_docx(worksheet_bytes)
        existing_ms_text = extract_docx(markscheme_bytes) if markscheme_bytes else ""
        step += 1

        # ── Step 1: Improve worksheet ─────────────────────────────────────────────
        yield f'data: {json.dumps({"step": step, "label": "Enhancing worksheet", "detail": "Improving quality"})}\n\n'
        improved_ws = await improve_worksheet(worksheet_text, spec_text=spec_text)
        step += 1

        # ── Step 2: Build / improve mark scheme ───────────────────────────────────
        if existing_ms_text.strip():
            # Teacher uploaded their own mark scheme — improve it rather than discard it
            yield f'data: {json.dumps({"step": step, "label": "Improving mark scheme", "detail": "Polishing uploaded mark scheme"})}\n\n'
            improved_ms = await improve_markscheme(improved_ws, existing_ms_text, spec_text=spec_text)
        else:
            yield f'data: {json.dumps({"step": step, "label": "Generating mark scheme", "detail": "Creating marking points"})}\n\n'
            improved_ms = await generate_markscheme(improved_ws, spec_text=spec_text)
        step += 1

        # ── Steps 3-6: Agents 1-4 run IN PARALLEL ────────────────────────────────
        yield f'data: {json.dumps({"step": step, "label": "Agents 1–4", "detail": "Running all checks in parallel"})}\n\n'
        combined = f"WORKSHEET:\n{improved_ws}\n\nMARK SCHEME:\n{improved_ms}"
        combined_input = f"WORKSHEET AND MARK SCHEME:\n{combined}"

        if spec_text.strip():
            # Spec provided — run all 4 agents including topic coverage
            coverage_input = f"{combined_input}\n\nINTENDED SCOPE:\n{spec_text}"
            r1, r2, r3, r4 = await asyncio.gather(
                run_agent(AGENT_1_PROMPT, combined_input),
                run_agent(AGENT_2_PROMPT, combined_input),
                run_agent(AGENT_3_PROMPT, combined_input),
                run_agent(AGENT_4_PROMPT, coverage_input),
            )
        else:
            # No spec — skip Agent 4 (topic coverage is meaningless without a spec)
            r1, r2, r3 = await asyncio.gather(
                run_agent(AGENT_1_PROMPT, combined_input),
                run_agent(AGENT_2_PROMPT, combined_input),
                run_agent(AGENT_3_PROMPT, combined_input),
            )
            r4 = "No specification provided — topic coverage evaluation skipped."

        step += 4  # accounts for steps 3, 4, 5, 6

        # ── Step 7: Agent 5 — final intelligent revision (gpt-4o) ─────────────────
        yield f'data: {json.dumps({"step": step, "label": "Agent 5", "detail": "Finalising revision"})}\n\n'
        agent5_input = (
            f"ORIGINAL WORKSHEET:\n{improved_ws}\n\nORIGINAL MARK SCHEME:\n{improved_ms}\n\n"
            f"INTENDED SCOPE:\n{spec_text}\n\n"
            f"AGENT 1 REPORT:\n{r1}\n\nAGENT 2 REPORT:\n{r2}\n\n"
            f"AGENT 3 REPORT:\n{r3}\n\nAGENT 4 REPORT:\n{r4}"
        )
        # gpt-4o for Agent 5 — the critical rewrite step that synthesises all reports
        final_text = await run_agent(AGENT_5_PROMPT, agent5_input, model="gpt-4o")
        revised_ws, revised_ms = parse_revised_output(final_text)
        if revised_ws:
            improved_ws = revised_ws
        if revised_ms:
            improved_ms = revised_ms

        # Tally issues found by the four agents
        pipeline_stats = _count_agent_issues(r1, r2, r3, r4)

        # Final event with results
        yield f'data: {json.dumps({"done": True, "worksheet": improved_ws, "markscheme": improved_ms, "pipeline_stats": pipeline_stats})}\n\n'

    except Exception as exc:
        # Surface the error to the frontend instead of dying silently
        yield f'data: {json.dumps({"error": True, "message": str(exc)})}\n\n'


@app.post("/api/process")
async def process_worksheet(
    worksheet: UploadFile = File(...),
    markscheme: Optional[UploadFile] = File(None),
    spec_txt: Optional[UploadFile] = File(None),
    spec_docx: Optional[UploadFile] = File(None),
    pasted_spec: Optional[str] = Form(None),
):
    """
    Process a worksheet through the full validation pipeline.
    Returns Server-Sent Events stream with progress updates.
    """
    worksheet_bytes = await worksheet.read()
    markscheme_bytes = await markscheme.read() if markscheme else None
    spec_txt_bytes = await spec_txt.read() if spec_txt else None
    spec_docx_bytes = await spec_docx.read() if spec_docx else None

    spec_text = read_spec_text(spec_txt_bytes, spec_docx_bytes, pasted_spec)

    return StreamingResponse(
        process_stream_generator(worksheet_bytes, markscheme_bytes, spec_text),
        media_type="text/event-stream",
    )


@app.post("/api/export/worksheet")
async def export_worksheet(request: Request):
    """Export worksheet text as a DOCX file. Result is cached by content hash."""
    body = await request.json()
    text = body.get("text", "")
    filename = body.get("filename", "worksheet.docx")

    cache_key = "ws_" + hashlib.sha256(text.encode()).hexdigest()
    if cache_key not in _export_cache:
        if len(_export_cache) >= 100:
            # Evict oldest 50 entries
            for k in list(_export_cache.keys())[:50]:
                del _export_cache[k]
        spec = await run_formatting_agent(text)
        docx_bytes = build_formatted_docx(spec)
        _export_cache[cache_key] = docx_bytes.getvalue()

    return StreamingResponse(
        iter([_export_cache[cache_key]]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post("/api/export/markscheme")
async def export_markscheme(request: Request):
    """Export mark scheme text as a DOCX file. Result is cached by content hash."""
    body = await request.json()
    text = body.get("text", "")
    filename = body.get("filename", "markscheme.docx")

    cache_key = "ms_" + hashlib.sha256(text.encode()).hexdigest()
    if cache_key not in _export_cache:
        if len(_export_cache) >= 100:
            for k in list(_export_cache.keys())[:50]:
                del _export_cache[k]
        docx_bytes = build_markscheme_docx(text)
        _export_cache[cache_key] = docx_bytes.getvalue()

    return StreamingResponse(
        iter([_export_cache[cache_key]]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post("/api/chat")
async def chat_endpoint(request: Request):
    """
    Chat endpoint for interactive revisions.
    Body: {"message": str, "worksheet": str, "markscheme": str, "history": list}
    Returns: {"reply": str, "updated_worksheet": str|null, "updated_markscheme": str|null}
    """
    body = await request.json()
    message = body.get("message", "")
    worksheet = body.get("worksheet", "")
    markscheme = body.get("markscheme", "")
    spec = body.get("spec", "").strip()
    # History: list of {role, text} — cap at last 10 messages to control token usage
    history = body.get("history", [])[-10:]

    spec_rule = ""
    if spec:
        spec_rule = f"""
SPECIFICATION CONTENT — APPLIES TO EVERY EDIT:
{spec}

CRITICAL SPEC RULE: Every scientific fact, definition, equation, and value you introduce
must be present in the specification above. Do NOT add content outside this specification,
even if it seems correct or relevant. If an instruction would require introducing off-spec
content, explain why and suggest a spec-compliant alternative instead.
"""

    system_prompt = f"""You are an AI assistant helping a teacher edit a GCSE science worksheet and mark scheme.
The worksheet may cover any GCSE science subject: Biology, Chemistry, Physics, or combined science.
Apply the same exam standards regardless of the specific science subject.
{spec_rule}
You have access to the conversation history so you can refer back to earlier requests (e.g. "undo that", "change it to 5 marks instead").

Rules:
1. If the instruction requires editing the worksheet, return the full updated worksheet after --- REVISED WORKSHEET ---.
2. If the instruction requires editing the mark scheme, return the full updated mark scheme after --- REVISED MARK SCHEME ---.
3. If only one document needs updating, only include that section.
4. Always start with a brief plain-English explanation of what you changed (2-3 sentences max).
5. Preserve all formatting, question numbers, and answer lines exactly.
6. Do not add commentary inside the document text itself.

Format your response as:
[Your brief explanation here]

--- REVISED WORKSHEET ---
[full updated worksheet, if changed]

--- REVISED MARK SCHEME ---
[full updated mark scheme, if changed]
"""

    # Build conversation: system + history (text only) + current message with full doc context
    messages = [{"role": "system", "content": system_prompt}]
    for msg in history:
        messages.append({"role": msg["role"], "content": msg["text"]})
    messages.append({"role": "user", "content": f"""INSTRUCTION: {message}

CURRENT WORKSHEET:
{worksheet}

CURRENT MARK SCHEME:
{markscheme}"""})

    response = await async_client.chat.completions.create(
        model="gpt-4o-mini",
        messages=messages,
        temperature=0,
    )
    full_reply = response.choices[0].message.content

    # Extract explanation (everything before first --- marker)
    ws_marker = re.search(r"-{2,}\s*REVISED\s+WORKSHEET\s*-{2,}", full_reply, re.IGNORECASE)
    ms_marker = re.search(r"-{2,}\s*REVISED\s+MARK\s+SCHEME\s*-{2,}", full_reply, re.IGNORECASE)

    explanation = full_reply
    updated_ws = None
    updated_ms = None

    if ws_marker:
        explanation = full_reply[: ws_marker.start()].strip()
        ws_end = ms_marker.start() if ms_marker else len(full_reply)
        updated_ws = full_reply[ws_marker.end() : ws_end].strip()

    if ms_marker:
        if not ws_marker:
            explanation = full_reply[: ms_marker.start()].strip()
        updated_ms = full_reply[ms_marker.end() :].strip()

    return {
        "reply": explanation or "Done — documents updated.",
        "updated_worksheet": updated_ws,
        "updated_markscheme": updated_ms,
    }


@app.post("/api/spec-ocr")
async def spec_ocr(file: UploadFile = File(...)):
    """
    Extract specification text from an uploaded image (e.g. a photo of a CGP book page).
    Uses GPT-4o vision to accurately read and structure the content.
    Accepts: image/jpeg, image/png, image/webp, image/gif
    Returns: {"text": str}
    """
    contents = await file.read()

    # Detect MIME type from filename extension; default to jpeg
    ext = (file.filename or "").rsplit(".", 1)[-1].lower()
    mime_map = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
                "webp": "image/webp", "gif": "image/gif"}
    mime = mime_map.get(ext, "image/jpeg")

    b64 = base64.b64encode(contents).decode()

    response = await async_client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            "This is a page from a GCSE specification document or revision book. "
                            "Extract ALL the scientific content accurately: topic headings, subtopics, "
                            "bullet points, learning objectives, equations, definitions, key facts, and values. "
                            "\n\n"
                            "IMPORTANT — DO NOT INCLUDE:\n"
                            "- Topic reference numbers or codes (e.g. '4.1.2', 'P1', 'B2.3', 'C3a', '3.4.1')\n"
                            "- Chapter numbers, section numbers, or page numbers\n"
                            "- Exam board codes or specification reference codes\n"
                            "- Any purely numerical or alphanumerical label that identifies a section rather than describes content\n"
                            "\n"
                            "DO INCLUDE everything else: all scientific facts, definitions, equations, "
                            "topic names (as plain words), and any content a student would need to know.\n"
                            "\n"
                            "Output clean plain text only — no markdown, no bullet symbols, no extra commentary. "
                            "If there are multiple columns, read them left-to-right. "
                            "Preserve line breaks for separate points."
                        ),
                    },
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:{mime};base64,{b64}",
                            "detail": "high",
                        },
                    },
                ],
            }
        ],
        max_tokens=4096,
    )
    extracted = response.choices[0].message.content.strip()
    return {"text": extracted}


@app.get("/logo.png")
async def serve_logo():
    """Serve the examqa logo."""
    logo_file = Path(__file__).parent.parent / "logo.png"
    if logo_file.exists():
        return FileResponse(str(logo_file), media_type="image/png")
    return {"error": "Logo not found"}


@app.get("/")
async def serve_frontend():
    """Serve the frontend HTML file."""
    frontend_file = Path(__file__).parent.parent / "frontend" / "index.html"
    if frontend_file.exists():
        return FileResponse(str(frontend_file), media_type="text/html")
    return {"message": "Frontend not found"}


# ============================================================================
# MAIN
# ============================================================================

if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8000)
