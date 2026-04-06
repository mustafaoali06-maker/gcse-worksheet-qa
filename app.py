import streamlit as st
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_TAB_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openai import OpenAI
import os
import re
import time
import json
import base64
from io import BytesIO
from agents import (
    FORMATTING_AGENT_PROMPT,
    AGENT_1_PROMPT,
    AGENT_2_PROMPT,
    AGENT_3_PROMPT,
    AGENT_4_PROMPT,
    AGENT_5_PROMPT,
)

# ---------------- CONFIG ----------------
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
st.set_page_config(page_title="GCSE Worksheet QA Studio", layout="wide")

ANSWER_LINE = "____________________________________________________________________________"

# ---------------- DARK MODE TOGGLE ----------------
if "dark_mode" not in st.session_state:
    st.session_state["dark_mode"] = False

# ---------------- THEME VARIABLES (Figma design tokens) ----------------
_DM = st.session_state["dark_mode"]

if _DM:
    # Dark theme
    _bg          = "#0f1117"
    _surface     = "#1a1f2e"
    _surface2    = "#1e2535"
    _border      = "#2a3248"
    _border2     = "#374151"
    _text        = "#e8ecf4"
    _text_mid    = "#9ca3af"
    _text_muted  = "#8b95a8"
    _text_light  = "#6b7280"
    _primary     = "#3b82f6"
    _primary_lt  = "#1e3a5f"
    _primary_dk  = "#2563eb"
    _success     = "#22c55e"
    _success_lt  = "#052e16"
    _error       = "#ef4444"
    _violet      = "#a78bfa"
    _violet_lt   = "#2d1f4e"
    _card_bg     = "#1a1f2e"
    _code_bg     = "#161b27"
    _shadow      = "0 4px 24px rgba(0,0,0,0.45)"
    _shadow_sm   = "0 2px 8px rgba(0,0,0,0.3)"
    _nav_bg      = "#1a1f2e"
    _sb_bg       = "#161b27"
else:
    # Light theme (Figma design)
    _bg          = "#f8f9fc"
    _surface     = "#ffffff"
    _surface2    = "#eef1f8"
    _border      = "#e5e8ee"
    _border2     = "#d1d5db"
    _text        = "#111827"
    _text_mid    = "#4e5c73"
    _text_muted  = "#6b7280"
    _text_light  = "#9eaabb"
    _primary     = "#2563eb"
    _primary_lt  = "#eff4ff"
    _primary_dk  = "#1d4ed8"
    _success     = "#16a34a"
    _success_lt  = "#ecfdf5"
    _error       = "#dc2626"
    _violet      = "#7c3aed"
    _violet_lt   = "#f5f3ff"
    _card_bg     = "#ffffff"
    _code_bg     = "#f5f6fa"
    _shadow      = "0 1px 3px rgba(0,0,0,0.08), 0 4px 16px rgba(0,0,0,0.06)"
    _shadow_sm   = "0 1px 2px rgba(0,0,0,0.05), 0 2px 6px rgba(0,0,0,0.04)"
    _nav_bg      = "#ffffff"
    _sb_bg       = "#fafbfe"

# ---------------- STYLE (matches Figma design) ----------------
st.markdown(f"""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');

html, body, [class*="css"] {{
    background-color: {_bg};
    color: {_text};
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
    -webkit-font-smoothing: antialiased;
}}

/* ── Sidebar ── */
section[data-testid="stSidebar"] {{
    background: {_sb_bg};
    border-right: 1px solid {_border};
}}
section[data-testid="stSidebar"] * {{ color: {_text} !important; }}

/* ── Sidebar logo zone ── */
.sb-logo-zone {{
    padding: 18px 20px 14px 20px;
    border-bottom: 1px solid {_border};
    margin-bottom: 0;
}}
.sb-logo-name {{
    font-size: 1.05rem;
    font-weight: 800;
    color: {_primary};
    letter-spacing: -0.01em;
}}
.sb-logo-dot {{
    color: {_primary};
    font-weight: 900;
}}
.sb-logo-sub {{
    font-size: 0.68rem;
    font-weight: 600;
    color: {_text_light};
    letter-spacing: 0.1em;
    text-transform: uppercase;
    margin-top: 2px;
}}
.sb-section-label {{
    font-size: 0.66rem;
    font-weight: 700;
    color: {_text_light};
    letter-spacing: 0.1em;
    text-transform: uppercase;
    margin: 16px 20px 6px 20px;
}}
.sb-divider {{
    height: 1px;
    background: {_border};
    margin: 14px 20px;
}}

/* ── Upload zones ── */
.upload-primary {{
    border: 1.5px dashed {_primary};
    background: {_primary_lt};
    border-radius: 8px;
    padding: 10px 12px;
    margin: 4px 0 10px 0;
    font-size: 12px;
    color: {_primary};
    font-weight: 600;
}}
.upload-secondary {{
    border: 1.5px dashed {_border2};
    background: {_bg};
    border-radius: 8px;
    padding: 10px 12px;
    margin: 4px 0 10px 0;
    font-size: 12px;
    color: {_text_muted};
    font-weight: 500;
}}
.upload-hint {{
    font-size: 10px;
    font-weight: 400;
    color: {_text_light};
    margin-top: 1px;
}}

/* ── Navbar / top bar ── */
.app-navbar {{
    background: {_nav_bg};
    border-bottom: 1px solid {_border};
    padding: 14px 32px;
    display: flex;
    align-items: center;
    gap: 16px;
    margin-bottom: 0;
}}
.app-navbar-title {{
    font-size: 0.95rem;
    font-weight: 700;
    color: {_text};
    flex: 1;
}}
.app-navbar-sub {{
    font-size: 0.72rem;
    color: {_text_muted};
    margin-top: 1px;
}}
.theme-pill {{
    background: {_surface2};
    border: 1px solid {_border};
    border-radius: 20px;
    padding: 5px 12px;
    font-size: 0.75rem;
    color: {_text_mid};
    font-weight: 500;
    white-space: nowrap;
    cursor: pointer;
}}

/* ── Buttons ── */
.stButton>button {{
    background: {_primary};
    color: #ffffff;
    font-weight: 600;
    font-size: 0.85rem;
    border: none;
    border-radius: 8px;
    padding: 8px 18px;
    transition: all 0.18s ease;
    box-shadow: 0 1px 2px rgba(37,99,235,0.2), 0 4px 12px rgba(37,99,235,0.15);
    letter-spacing: 0.01em;
    font-family: 'Inter', sans-serif;
}}
.stButton>button:hover {{
    background: {_primary_dk};
    box-shadow: 0 2px 4px rgba(37,99,235,0.3), 0 6px 20px rgba(37,99,235,0.2);
    transform: translateY(-1px);
}}
.stButton>button:active {{
    transform: translateY(0);
}}
.stDownloadButton>button {{
    background: {_success};
    color: white;
    font-weight: 600;
    border: none;
    border-radius: 8px;
    font-family: 'Inter', sans-serif;
    box-shadow: 0 1px 2px rgba(22,163,74,0.2), 0 4px 12px rgba(22,163,74,0.12);
}}
.stDownloadButton>button:hover {{
    background: #15803d;
    box-shadow: 0 2px 4px rgba(22,163,74,0.3), 0 6px 20px rgba(22,163,74,0.18);
    transform: translateY(-1px);
}}

/* ── Headings ── */
h1, h2, h3 {{ color: {_text}; font-weight: 700; }}

/* ── Content cards ── */
.qa-card {{
    background: {_card_bg};
    border: 1px solid {_border};
    border-radius: 12px;
    padding: 18px 22px;
    margin: 10px 0;
    box-shadow: {_shadow_sm};
}}

/* ── Section row headers ── */
.section-row {{
    display: flex;
    align-items: center;
    gap: 10px;
    margin: 28px 0 8px 0;
}}
.section-bar {{
    width: 3px;
    height: 22px;
    border-radius: 2px;
    flex-shrink: 0;
}}
.section-title {{
    font-size: 0.95rem;
    font-weight: 700;
    color: {_text};
}}
.section-sub {{
    font-size: 0.72rem;
    color: {_text_muted};
    margin-top: 1px;
}}
.section-label {{
    font-size: 0.66rem;
    font-weight: 700;
    color: {_text_light};
    text-transform: uppercase;
    letter-spacing: 0.1em;
    margin-bottom: 4px;
}}

/* ── Text areas ── */
.stTextArea textarea {{
    background: {_code_bg};
    border: 1px solid {_border};
    border-radius: 8px;
    color: {_text};
    font-family: 'Fira Code', 'Courier New', monospace;
    font-size: 12.5px;
    line-height: 1.6;
}}
.stTextInput input {{
    background: {_surface};
    border: 1px solid {_border};
    border-radius: 8px;
    color: {_text};
    font-family: 'Inter', sans-serif;
}}

/* ── Dividers ── */
hr {{ border: none; border-top: 1px solid {_border}; opacity: 0.8; margin: 24px 0; }}

/* ── Expanders ── */
.streamlit-expanderHeader {{
    background: {_surface2};
    border-radius: 8px;
    font-weight: 600;
    color: {_text} !important;
    border: 1px solid {_border};
}}

/* ── Alert boxes ── */
.stAlert {{ border-radius: 8px !important; }}

/* ── Chat bubbles ── */
.chat-user {{
    background: {_primary_lt};
    border-left: 3px solid {_primary};
    border-radius: 0 8px 8px 0;
    padding: 10px 14px;
    margin: 8px 0 4px 0;
    font-size: 13.5px;
    color: {_text};
    line-height: 1.5;
}}
.chat-assistant {{
    background: {_card_bg};
    border: 1px solid {_border};
    border-left: 3px solid {_violet};
    border-radius: 0 8px 8px 0;
    padding: 10px 14px;
    margin: 4px 0 8px 0;
    font-size: 13.5px;
    color: {_text};
    line-height: 1.5;
}}
.chat-name {{
    font-size: 0.66rem;
    font-weight: 700;
    text-transform: uppercase;
    letter-spacing: 0.1em;
    margin-bottom: 4px;
}}
.chat-prompt-chips {{
    display: flex;
    flex-wrap: wrap;
    gap: 6px;
    margin-bottom: 12px;
}}
.chat-chip {{
    background: {_primary_lt};
    border: 1px solid rgba(37,99,235,0.2);
    border-radius: 20px;
    padding: 4px 12px;
    font-size: 11px;
    color: {_primary};
    font-weight: 500;
    cursor: pointer;
}}

/* ── Pipeline cards ── */
.pipeline-wrap {{
    background: {_card_bg};
    border: 1px solid {_border};
    border-radius: 12px;
    padding: 20px 24px;
    margin: 8px 0 16px 0;
    box-shadow: {_shadow};
}}
.pipeline-title {{
    font-size: 0.9rem;
    font-weight: 700;
    color: {_text};
    margin-bottom: 16px;
    display: flex;
    align-items: center;
    gap: 8px;
}}
.cards-grid {{
    display: grid;
    grid-template-columns: repeat(8, 1fr);
    gap: 8px;
}}
.agent-card {{
    border-radius: 10px;
    padding: 12px 8px;
    text-align: center;
    border: 1px solid {_border};
    background: {_surface2};
    transition: all 0.25s ease;
    position: relative;
    overflow: hidden;
}}
.agent-card.done {{
    border-color: {_success};
    background: {_success_lt};
}}
.agent-card.active {{
    border-color: {_primary};
    background: {_primary_lt};
    box-shadow: 0 0 0 2px rgba(37,99,235,0.15);
    animation: cardPulse 2s ease-in-out infinite;
}}
.agent-card.pending {{
    opacity: 0.42;
}}
@keyframes cardPulse {{
    0%   {{ box-shadow: 0 0 0 2px rgba(37,99,235,0.15); }}
    50%  {{ box-shadow: 0 0 0 6px rgba(37,99,235,0.0); }}
    100% {{ box-shadow: 0 0 0 2px rgba(37,99,235,0.15); }}
}}
.card-icon {{ font-size: 1.4rem; line-height: 1; margin-bottom: 5px; }}
.card-label {{
    font-size: 0.62rem;
    font-weight: 600;
    color: {_text};
    line-height: 1.3;
    margin-bottom: 5px;
}}
.agent-card.done .card-label  {{ color: {_success}; }}
.agent-card.active .card-label {{ color: {_primary}; }}
.card-status-dot {{
    width: 6px; height: 6px;
    border-radius: 50%;
    margin: 0 auto;
    background: {_border2};
}}
.agent-card.done .card-status-dot {{ background: {_success}; }}
.agent-card.active .card-status-dot {{
    background: {_primary};
    animation: dotBlink 1s ease-in-out infinite;
}}
@keyframes dotBlink {{
    0%, 100% {{ opacity: 1; transform: scale(1); }}
    50%       {{ opacity: 0.3; transform: scale(0.75); }}
}}
.card-check {{
    position: absolute;
    top: 5px; right: 7px;
    font-size: 0.65rem;
    color: {_success};
    font-weight: 700;
}}
.pipeline-progress {{
    margin-top: 14px;
    height: 3px;
    background: {_border};
    border-radius: 3px;
    overflow: hidden;
}}
.pipeline-progress-bar {{
    height: 100%;
    background: linear-gradient(90deg, {_primary} 0%, {_success} 100%);
    border-radius: 3px;
    transition: width 0.6s ease;
}}
.pipeline-step-label {{
    margin-top: 8px;
    font-size: 0.72rem;
    color: {_text_muted};
}}

/* ── Export cards ── */
.export-card {{
    background: {_card_bg};
    border: 1px solid {_border};
    border-radius: 12px;
    padding: 0;
    overflow: hidden;
    box-shadow: {_shadow_sm};
}}
.export-card-stripe {{
    height: 4px;
    width: 100%;
}}
.export-card-body {{
    padding: 14px 18px 16px 18px;
}}
.export-card-title {{
    font-size: 0.85rem;
    font-weight: 700;
    color: {_text};
    margin-bottom: 4px;
}}
.export-card-desc {{
    font-size: 0.72rem;
    color: {_text_muted};
    line-height: 1.5;
    margin-bottom: 12px;
}}
</style>
""", unsafe_allow_html=True)

# ---------------- HEADER ----------------
_LOGO_B64 = "data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAPIAAADQCAMAAAAK0syrAAAA5FBMVEX///8ANF8ZTrcLQ7MAMl4AMF0AKFgAJFYAJlcALVv5+/0AKlkOOmMAIlUAHlMALFoAG1IlRGkxUHIAGVEAP7MAFU8AQbMAN7AARrUAPbLz9ffi5urr7vEPSrYAM68mSW6lrruHlqgAD02XpLPS2N6zvMdoe5NPZoJGX310jM1Pb8IlUbdecovY3eOAj6K/xs+eqrjP2O62weOisdxzhZrGzdWOmqs6VnZkdo7Czunl6vZWeMY7Yb4ALa6GmtOTpNbc4vIAAERkgspOccMxWrsAAEe8xuWot98AIax6kc+NntQ1YL7SU2t8AAASVElEQVR4nO1caVviShYGs6+EHUKQICAuCKgs7dKt0+3o9L3///9MllpOJQExMI/Xnno/aSWp1Kk6+zmhUODg4ODg4ODg4ODg4ODg4ODg4ODg4ODg4ODg4ODg4ODg4ODg4ODg4ODg4ODg4ODg4OD4JOiuq+d7LvOxYLr91pOENxosJuPxeLJYjTr7TzccTCqmKhfL68l04C+9XWl3R4t1UTbPpkN2fNkflw2tPF5d7b+2CFf9imOpRknTtJKhWvXelKF60Ktg9FbwwpqMV8pgLd6q7KglSSwWRVHSZFWx68X1pD8aeon3jsrkeTM8Qm+h2UbwnCjJ9TO6BK9vOrIWTqep1UlykjxYjuuyVIQI3rgAXLSsShilHnhwapHx6oAOryxVLCYQkm45tjnp+0N64lOVTODoBb1vy/TBkjWLb3KnDpzPMPY+aG9c15ILDCBLgLUM8krRolsxrNO7J2T0qqxkTIcf12TlX0tyb5nstFgKHrTYe6vRCnxZZueQ5D3PeaQY2auTbErzKb0H730Ak2yEVCYnN6pLmdNRGGTTPJW+bTyrJh+UTL3gnTkpjjHGe1Hcr6dmJG8UydpmNhmV+3hwQRdsE14b1LOmgtDOyMuXdNpiWU3fao+GRiljivo+rL1wtixOXuDbdLoesYzXS4lzfHyjf/IexUV5St6+yiATQizLmQdikIV9HCs7a0YClQjNmIp7PR7UKVtTQV5u5hkCIBln78mAmD2daOam2K8nXpB4heLjO0dUs6AlnxKlIuFzL3jZh8K+BOg/6f3bs1HP6zl0mDPWLLvYKxs2WIZG9IRnkeGYq2ZUIhzy/jGjCSXFqTuWYmgMYSK1csMkj4klRcmyHkVNURiZBpzyMazB9GJ1PIq2f0kNB5XbgAfJvRFXufSA6iN8z4hRDPba73id2XwxlqpgK4AczhOibFinvj9hLVU07kzm/ikct/x8FPtgkyWNGMsOOGeNMOGcGtt6MDghRKin+BYg3SH/0lXpnSkwcmSHoOkLn3C+dRJvQuOn8TigWQGOz0dA/YvAAgLrDlSVTYY7Crk7WPSI7JZE+XQA1iQqS/iqBSUOGBiG4zUJM6vJCIJh4pnKdFxl3N6dAXfTguoArK9Kx3uE342pS9WUTe7Qoe6zR/BV4GHRJE4Lo0rUNdndBTx8a0wYDRyFmu+Ui3SF1hxeAPzm0MMfEMET1xOiS6qUNB8csswaTqD8St/I6Aw8YFGHtdAHDqZFjTjcCmpKPgLgUFErE2FN9RfwZoF6Fcl2q4C0ClB7IhvZAhMHzgeQwPgWwEGhrl6AKd0Ki2WiHUFPKrFnQ6p3ARcyooShVcBjVTqe1C5AeznDrAltKFjglB3oWU43jO8KF3ghMjwTDxgpaZ29brwjKlB6YEGilEhfAC+rRDZRh1sL7z4lhyEqMGYCh+TkybYAQdKAHHkDAziBwB1mYwD0YugQAD2bdIE9qijBJgLJMk7h7VSwxKKbPc5K4o4AOqq0mI1C+KtFr84E9w7MxuhmgrMVuCFX9obH2M0C1mVF2UKB6tOlhyHBfIRL9a2WK3qEklmyIiiqkciNVJhHWM+hWIJcXxjQk0w5/QMQZVJzPQHmH4rmFWV4oN6ZcUar7Qrv/SAv8BpYT3bGcLYoMbkJYDRToR0wASf0oRMwFRRNoN6Z0x/RTbXzuNijtCebgpVYusvwQJVxryDXJC0I4EggysAwgKRBgbFdzExAP1bzBFKDd4LzcCtTAgPDW5XlrQ6IKOqJ1BTgSMAA4NBYXwrwhANJo5sqSjkoZp26LGj1lLxA9yqpMwHTbxNlEGt8oybHhgzjAvUOdYlHD7nEKPhdsT0fIcr18TD5yJCRhYQz4NOFptTpJJVRCQH0PzTvcPMY7QXVPuMf74qsJFpMrWYojjlNezeuyexSwrEHFsdIngElTiySQY+6QqxhgLbLB+PQ92LVyG5wgcIWVSWApViWbTtOebyYZyqHM1YUGJPJCIqcCOyAxQZOD9CfrIYHqp8RZSDidg6K4R6L0sAPEfgis+XVxtrRNKniGbcYLjSpsIEKAMe2KUYA6RaxqMNxutuMgt+ZZBoD7Oa8jVLJX/YwwRkk8lIdEKQCt4wqE5EJ1iFPQNKAKBvQ69sZHUpywsXacL+VCqOYkANGjqyf4JkgwVYnxwaiGnbP/cw4k/Hu7DyiDM3oLiTroHZEjwyquMqGU3bLgGKgy8GhlUBUw4RLDGnpRPpHSQaM3Xv/9gnNWVNRY4JiYPQUIJrMGUPXGNhqNloHKSMHhFE6CE2LhTwAJLNWMRMDciTa2ZxSD/cKOBYg4uxITFYa6CmYUoQeAEgZMVoKiHjO6gzQ2Gk7msSSWjTLu6KbBdUOyGRQSZlZrMMDOBLoZSafsMl2gdxkvrxXQQenzKTUMDy6Do+G0GG5DZSigM4G3lfRRmtaJIuQBrkdcBnrrGUnvFlrnbPsyDicqjm4cgu663WGy9G8v5iUT/5Fua1C3hbVdSkLQ8XHpEycxdD15mKiEg51PDg0NvaFMQUUOKA25XwUJxJZomJjWKocFpGo6qBV5JgFAeuB/XaZ+eRgonTvBOBUkG1gbBqIHZjoBFrrvPX02fYqa1HBN/rUnFWjxXUou0KdfZZZPguOlv4JOLUH8sZQlDfZLiA3OQsVhURxJAXi+IDsNc5ZgygBOBF+Zs5BlIEEUdXc2RRTbLJdgCntVIi3K/pbcwTY0IDwibSDQA1DdbabuYf1EYgcq0Q4Z5tiCsArTM6QTg7L0x+EvvWYcURKC8YiidiAMEOdPcg45uqqAApVFt3vDTEFSKKKNogpoLO4g+e0CcttCT/kNPZB7wBRrMA9Zt5fSUqzWF8VPFqwBJFjZUOuZ5OWAlyRK7mJ4W/p6ziJdPGI9h/BUgwMIYCNTLhaRS30PK+yWomAemdKQFBLMaRtygB+GEszHR+FaRHFjrVXR5ZEDEhbXyHDjJ/N9LiJ9jo8v5Elxj0okmST9M2wSiYwYK6ncCrTF0LbVdbo+H79pe7KtNVApsNVhd2WJVlxjPJ4Pow0hF4pmhgazDYNZTJuMpLlLhQl6rQUS7bpR0OLE7UkFc1yr1fpkcBopWVPXOjRiVUYU9hkuLiHKCMs+2NTUi1L1Yrm2Wnfh42lOgDz0MYLBW8wLsqWYk5G6ILX8bxUi/KHJ978wpwIHM2Ol69zOgvBbIeaioPjcPDubs8vLrvd7svF+e3dgXv7/4G4f3ps1oRu9yhEtyXUmhe395+9qP8l7i7aQkwtRVdoP1x/9sL+V7h7rCXpRVTXLv5Iou8f2tkEx0Sf/3lCfdNqbSQ4ROvo7rOXeGD8bm8lOET76bMXeUjoD413KT46qp1/9joPB/2XsAPFR0eNn5+90oPhYTeKA5p3POeOP+j3B/AbMIQrf9VfDfytmWh3NpiuZjr7v3+IT+Aofu/C1THaf70/XWdarlqqLKtBJLruA7LdQXAhGFetqrlgQo+O2QtRXgR/Lkq2Kqu26MeXhhM5/N+yFweLewqFt/c1F0G3+fredNMq+KZEkh1SfxlpNCkkGnVYIvHtOOEwLazwJ4BidENnUseVArVyMJrvd+XqiORWYzuHdcxEyp6UKBaJvJMFqmwok2/0JyC9Xp8V5haoCCh7fBzF4tdmBySD5G7tYdtkV3ay0wingo8epBKhCy384m1aC2TPp7JRtXbAPJM83tQ9QHAYa37e4JBnfSaFq47eMTz9PcKbay/5oUEq0LVm5ep9S0F8+QnEUW71sno3mdUuB/go/YEa52QE+L1G1FMwIpAQDvzjQFLZ6J8rky6l3q8K74e1DhxwwtrDlmEn2XjZOB4Ppt3XR+bcfXiAfjijmYOZX8InXkWajrV6SM/b7QH5FRZuOxuhy7gIcixdGkluNWq3RShAZgfwrCLXHDXORTnObKpqrSM3i6nA1Tk5PUXEIN4OQupXaC3md1uKk+AFUgWUbcvLiGh5y++XH693d6+8jEEN2X36fB3joUpKF78/Zk+Eyl5UsKCydxAVUh0GqzcWH7KCtwkUdw4x9FlScOwzJ5/RIWy+EYW+OCM3CbTRy3wYkN7P9ER01Q2ipTjRUXKNLRtUHVOzDbcoWrgMgBa6dIcZHTb2HYexLQrEAjY/7iGluP8cjf3cpyY1sBYa/kKsm3ckr9EUh/Q4OFWKQPkKHSAvKvfh+UlYdMbfvB8rX3Yt45DlOdOl4L/5Gd94KlGThe2Y2DHVspc8C9S6CCyzJqB5LW+yR/iLdNGhP8hfTAd6I59WIeOjtpV27eA7/uo45ufUj+Dv0MZ/bgOTjLK8TF0rTdTLUMgAuIBpQbRn1lpOCMiq4084JtGf7VeAQiCgLUfj/1oyIis4wdsraYc4r4uNLQHLzR8Zc+NOwajJjhIuPCvWRkSzHHYu4O4j016ICNuUKJNvOIbwvYqLa4WxuHFG1oqj4NfrnMhz+T7gHP1qU5NqvjLnQQtP9kfgC0Ldn8NxTFK4YHiCd1GxdNi+I6Ykk+Q0FkbWQ/vtQzLthfPx6/FYgYh+RnKm/UGdAWscgtgQFYywDMSuz6rtArDj5KAFp9LQlyAEXk9yKzA42vrXIWkWHH/51HilzHTB2o5ExGVpoWscgRwJ0s3qI1WPvC4k6aVLVkfYin8lg7ZXzM20GJG6MrS+2TPS/Vriml6Nm+P95l5AstDMmQ5QpKR2DtBMoyiOFHcsA9tmIDkAf3Igi5mO0mfl6kjeR3AhZt4AjjFiXBSR3Q6l+Pu42w8N+rVGSmxmTIcpSP4xANDnNgyAaYlZH7iXVz35StlEjz0G014ZTJhsQ/XFz8esi3AM30l+tTSQTypIku7H7CHrr8Y8dxD1GWAcQvxzZd9KTgRplD+NuuthGtaJz/YlIjiyTXkN6jCC6HJOcIcuIsjRj4wuU5EVMJOqrRt4oNbrY98J8PE8o8L2gsxr7FfliQihEgYbGHhlCpNBjii/R0MyPsaRf+aVWRo4fyzL2P2N1hJ8jbir+7ANHljhLlLMnOQmSH2iH8+txNNGI1HdghwW2POEJWJRreC9MK/zeSDmZh98iiAmtg4EuYCOEP1oQTyKacExBmkaxGScmCe/JYdJAmJWRMF83W91uHA7rDRRS3IVHL0TSfdHFovw7fhy3YkbmBAe9xiTxEtTNLFoRp3YqyAYpsTVDbEv74FC/JPkBB+Sc7fGzQAxuiY8tRFv+fP54ER/tz0DMI4cj8rqQqyLEhywc38SPowROLJTkRxMMc+X788FqsY6dD9LQai18f4HTOvgLCbQh1JyvEx55yjnbD3epSArhqYZCCqTG2yG3PrcRyThHwIRIS9JIL8oBs6uyYcc6lyRLREMhP6cm1pF2MhOBIjHTmI8PK8pUfwUE/gKfFjx8b3VjxY1yA7UodnpsxNqrhe5D7j5yiyrJfC62V4uMpmAHmSD0ARZtUkXd0PQzmV4iztoXP0HKp3uLidafb84fjy7DvjLkdwfOdvDPX+3Y90KijFsx0Qktk6lZfHJuMbkZoo0daqysSNMezhfgRkcX+WLOoUr6rzD31Whd/Lh9xpdeji4DkIvhP5ivURcFTmlhc9Kvsr9yZWHeHBpsQ6vkEMnFMQWJOZBnRvj4wKIMsh/opFtC+9dNxGPX6WxvrK6FxhF6eGVFlST6Mwx9B5AmyfRzv87apr+dJdll6i1XNLZJVTfjgTrm42ncI6vm+s4xE0+pilS31v19p8TEIL4gIH0dWKzop15PzXIIk8ZOwzPbkg0jLBTKlQUMA0Zj2bYUVVUsa+3TYS+eomxi16wjxgPk0/sxuuEQGZEYblZFqtXunqfHEcUNIXxsU7R+Neqfnk4Ho2XKcfCWo/lgMJ99foPnU3ZxOU1xCxzyF8ffu1UeMcXbSlJfBO5FC9RfNkLA2FSq+Dq4qbVb3XeJ7hKKj794J5R+c3kcOFQhz7ZaW1r7WoTi5taC+j8d928/m82GQLGBaECw0N5UdPyH4/r26cfDi3DcbAtJpKmGBAvtl6/ayOndnF9+DyiuNVI0RxzejW1UlyE35OrHr0pxhPu7p4dH4fj4uImQPvIkjh8O2Hn1WdDvr1/fbm+fnm5vb348zzpzitpX19UZuDtqbia4cfzyJzah60/tZvZJN5qomeDPg/dX9zgl04328dHtl9Zb26HfPLQDXd4g5DaPGz/fbdv86nCvn35dNpqhMhcuH56u/wA1vRPc++fn5/v/F2o5ODg4ODg4ODg4ODg4ODg4ODg4ODg4ODg4ODg4ODg4ODg4ODg4ODg4ODg4OFL4L8bTjHBm1zuAAAAAAElFTkSuQmCC"
# ── NAVBAR (top bar with title + dark mode toggle) ──────────────
_theme_icon = "☀️  Light" if _DM else "🌙  Dark"
_nav_col1, _nav_col2 = st.columns([9, 1])
with _nav_col1:
    st.markdown(f"""
<div class="app-navbar">
    <img src="{_LOGO_B64}" alt="examqa.com" style="height:36px;object-fit:contain;opacity:0.9;" />
    <div style="flex:1;">
        <div class="app-navbar-title">GCSE Worksheet QA Studio</div>
        <div class="app-navbar-sub">AI-powered exam paper enhancement</div>
    </div>
</div>
""", unsafe_allow_html=True)
with _nav_col2:
    st.markdown(f"<div style='height:14px'></div>", unsafe_allow_html=True)
    if st.button(_theme_icon, key="theme_toggle"):
        st.session_state["dark_mode"] = not st.session_state["dark_mode"]
        st.rerun()

# ── SIDEBAR ─────────────────────────────────────────────────────
with st.sidebar:
    # Logo zone
    st.markdown(f"""
<div class="sb-logo-zone">
    <div class="sb-logo-name">examqa<span class="sb-logo-dot">.</span></div>
    <div class="sb-logo-sub">Worksheet QA Studio</div>
</div>
""", unsafe_allow_html=True)

    # Documents section
    st.markdown('<div class="sb-section-label">Documents</div>', unsafe_allow_html=True)
    st.markdown(f"""
<div class="upload-primary">
    📄  Worksheet (.docx)<br>
    <span class="upload-hint">Drop file or click to upload</span>
</div>""", unsafe_allow_html=True)
    worksheet_file = st.file_uploader("Worksheet", type=["docx"], label_visibility="collapsed")

    st.markdown(f"""
<div class="upload-secondary">
    📋  Mark Scheme (.docx)<br>
    <span class="upload-hint">Drop file or click to upload</span>
</div>""", unsafe_allow_html=True)
    markscheme_file = st.file_uploader("Mark Scheme", type=["docx"], label_visibility="collapsed")

    # Spec section
    st.markdown('<div class="sb-divider"></div>', unsafe_allow_html=True)
    st.markdown('<div class="sb-section-label">Specification (optional)</div>', unsafe_allow_html=True)
    st.markdown(f"""
<div class="upload-secondary">
    📑  .txt or .docx file<br>
    <span class="upload-hint">Aids topic coverage checks</span>
</div>""", unsafe_allow_html=True)
    spec_txt  = st.file_uploader("Spec .txt",  type=["txt"],  label_visibility="collapsed")
    spec_docx = st.file_uploader("Spec .docx", type=["docx"], label_visibility="collapsed")

    st.markdown(f'<div style="font-size:0.7rem;color:{_text_light};margin:4px 0 2px 0;">Or paste below</div>', unsafe_allow_html=True)
    pasted_spec = st.text_area("Paste spec", label_visibility="collapsed",
                               placeholder="Paste specification text here…", height=90)

    # Run button
    st.markdown('<div class="sb-divider"></div>', unsafe_allow_html=True)
    run_button = st.button("▶  Run Enhancement Pipeline", use_container_width=True)
    st.markdown(f'<div style="font-size:0.68rem;color:{_text_light};text-align:center;margin-top:6px;">8-step AI pipeline · ~30 seconds</div>', unsafe_allow_html=True)

# ---------------- HELPERS ----------------

def extract_docx(file):
    doc = Document(file)
    return "\n".join([p.text for p in doc.paragraphs])

def clean_text(text):
    return re.sub(r'[#*]+', '', text)

def add_answer_lines(text):
    lines = text.split("\n")
    output = []
    for line in lines:
        output.append(line)
        match = re.search(r"\((\d+)\)", line)
        if match:
            marks = int(match.group(1))
            for _ in range(min(marks, 4)):
                output.append(ANSWER_LINE)
    return "\n".join(output)

def extract_total(text):
    match = re.search(r"Total for paper\s*=\s*(\d+)", text)
    return int(match.group(1)) if match else None

def fractional_marks_present(text):
    return bool(re.search(r"\(\d+\.\d+\)", text))

def keyword_overlap(text1, text2):
    words1 = set(re.findall(r'\b[a-zA-Z]{5,}\b', text1.lower()))
    words2 = set(re.findall(r'\b[a-zA-Z]{5,}\b', text2.lower()))
    if not words1:
        return 0
    return round((len(words1 & words2) / len(words1)) * 100, 1)

def extract_question_numbers(text):
    nums = set()
    for m in re.finditer(r"^\s*(\d+)\s*(?=[.(A-Za-z])", text, re.MULTILINE):
        n = m.group(1)
        try:
            v = int(n)
        except ValueError:
            continue
        if v <= 50:
            nums.add(n)
    return sorted(nums, key=lambda x: int(x))

def strip_answer_lines(text):
    lines = text.split("\n")
    return "\n".join([ln for ln in lines if ANSWER_LINE.strip() not in ln.strip()])

def detect_question_structure(text):
    ROMAN_RE = re.compile(r"^\s*\((i{1,4}|iv|vi{0,3}|ix|xi{0,3}|x{1,3})\)\s", re.IGNORECASE)
    PART_RE = re.compile(r"^\s*\(([a-z])\)\s")
    structure = {}
    current_q = None
    current_part = None
    for line in text.split("\n"):
        if "Total for question" in line:
            continue
        m_main = re.match(r"^\s*(\d+)\s*(?=[.(A-Za-z])", line)
        if m_main:
            v = int(m_main.group(1))
            if v <= 50:
                current_q = m_main.group(1)
                current_part = None
                structure.setdefault(current_q, {"parts": {}})
            continue
        if current_q is None:
            continue
        m_roman = ROMAN_RE.match(line)
        if m_roman and current_part is not None:
            roman = m_roman.group(1).lower()
            structure[current_q]["parts"].setdefault(current_part, set())
            structure[current_q]["parts"][current_part].add(roman)
            continue
        m_part = PART_RE.match(line)
        if m_part:
            letter = m_part.group(1)
            structure[current_q]["parts"].setdefault(letter, set())
            current_part = letter
    result = []
    for qnum, info in structure.items():
        parts_list = []
        for letter in sorted(info["parts"].keys()):
            roman_set = sorted(info["parts"][letter])
            entry = {"letter": letter}
            if roman_set:
                entry["roman_subparts"] = roman_set
            parts_list.append(entry)
        result.append({"question_number": qnum, "parts": parts_list})
    return result

def read_spec_text(spec_txt_file, spec_docx_file, pasted_spec_text):
    parts = []
    if spec_txt_file is not None:
        try:
            parts.append(spec_txt_file.read().decode("utf-8"))
        except Exception:
            pass
    if spec_docx_file is not None:
        try:
            parts.append(extract_docx(spec_docx_file))
        except Exception:
            pass
    if pasted_spec_text:
        parts.append(pasted_spec_text)
    return "\n\n".join(p.strip() for p in parts if p and p.strip())

# ---------------- AI ----------------

def improve_worksheet(text):
    prompt = """
You are improving a GCSE worksheet to match a professional exam-standard format.
Follow these rules EXACTLY:

CONTENT RULES:
1. Make every question clear and unambiguous. Remove AI-sounding or strange wording.
2. Questions should replicate real GCSE exam questions in style and difficulty.
3. Ensure a MIX of question types: 1-mark recall, 2-3 mark describe, 3-4 mark explain, calculation questions.
4. Include at least one application-style question with a named scenario (e.g. "A student, Sarah, connects a circuit...").
5. Do NOT repeat questions or topics.
6. Ensure every multi-mark question has appropriate cognitive demand.
7. Numbers in question stems should be written as WORDS (e.g. "two" not "2"), EXCEPT for: physical values (e.g. "15 m/s"), equations, or units.

FORMATTING RULES:
8. Remove ALL topic headers (e.g. "Work and Energy Transfers", "Forces", "Section A").
9. Remove ALL formatting symbols: *, #, bullet points, dashes used as headers.
10. In any question context/stem text, replace " = " and " - " used as label separators with ": ".
    Example: "Work Done = Force x Distance" in a context line -> "Work done: force x distance"
11. Question numbering must be consistent: 1, 2, 3 ... (a), (b), (c) ... (i), (ii), (iii).
    - Main question numbers should NOT have a dot (use "1" not "1.")
12. Do NOT add answer lines - these are handled separately.
13. Keep mark allocations exactly as shown, e.g. (2).
14. Ensure there is NO space between sub-parts (a), (b), (c) of the SAME question.
15. There SHOULD be a blank line between separate main questions (1, 2, 3...).
16. Do NOT completely rewrite questions - only improve clarity and GCSE realism.
17. If a question has sub-parts (a)(i), (a)(ii), the letter (a) alone should NOT be on its own line
    if it only introduces roman-numeral sub-parts. Use the format:
    (a) (i) question text here   (1)
        (ii) question text here  (2)

OUTPUT:
Return the improved worksheet only. No commentary or explanations.
"""
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": prompt},
            {"role": "user", "content": text}
        ],
        temperature=0
    )
    return add_answer_lines(clean_text(response.choices[0].message.content))


def generate_markscheme(text, mismatch_info: str = None):
    mismatch_block = ""
    if mismatch_info:
        mismatch_block = f"""
CRITICAL - SPECIFIC ISSUES TO FIX IN THIS REGENERATION:
{mismatch_info}
You MUST resolve every issue listed above. Do not reproduce these errors.
"""
    prompt = f"""
You are generating a fully explicit GCSE-style mark scheme from a worksheet.
{mismatch_block}

CONTENT RULES:
1. Write a marking entry for EVERY question and EVERY sub-part in the worksheet.
2. NEVER skip any sub-question.
3. NEVER use vague placeholders like "Working step (1)" or "Answer (1)".
4. For CALCULATION questions:
   - Show the actual equation, numerical substitution, and final answer with units.
   - Do NOT award a mark for giving the equation alone (guideline rule).
   - Award marks for correct substitution (1) and correct answer with units (1).
   - Example: "a = (v - u) / t = (0 - 20) / 5 = -4 m/s2 (1)(1)"
5. For NON-CALCULATION questions:
   - Give clear, specific marking points.
   - Allow reasonable alternatives: "OR" / "Accept..."
   - If 3+ possible answers exist, use: "Any [one/two/three] from:" followed by bullet points.
   - Use "OR" when only two alternatives exist.
6. Each mark MUST be a whole number - use (1) only. Never (2) for a single point.
7. Only the FIRST letter of each marking sentence should be capitalised.
8. Sentences longer than 3-4 words MUST end with a full stop before the (1).
9. Any useful side note should be prefixed with [NOTE]:
   e.g. "[NOTE]: Accept velocity instead of speed"

FORMATTING RULES:
10. Bold question numbers: "1", "2" etc. (just the number).
11. Sub-part labels in brackets: (a), (b), (c), (i), (ii).
12. NO space between marking points WITHIN the same question part.
13. A SMALL space (one blank line) between DIFFERENT sub-parts (a)->(b)->(c).
14. A SMALL space between the last mark of one question and the Total line.
15. Include a "(Total for question X is Y marks)" line after each main question.
    CRITICAL: Use "is" NOT "=" — e.g. "(Total for question 1 is 6 marks)" NOT "(Total for question 1 = 6 marks)"
16. At the very END of the mark scheme, add on its own line:
    "Total marks for question paper: Z"
    where Z is the sum of all question totals.

Question mapping and numbering:
- Use EXACTLY the question numbers and sub-parts from the DETECTED QUESTION STRUCTURE below.
- Do NOT invent new question numbers or sub-parts.
- Do NOT omit any question or sub-part.
- Treat numbers inside sentences (e.g. "2.0 m/s", "5 kg") as data NOT question numbers.
- Only start a new question number at the beginning of a line.
"""
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": prompt},
            {
                "role": "user",
                "content": (
                    "WORKSHEET TEXT:\n"
                    + text
                    + "\n\nDETECTED QUESTION STRUCTURE (do NOT change these IDs):\n"
                    + json.dumps(detect_question_structure(text), ensure_ascii=False)
                ),
            },
        ],
        temperature=0
    )
    return clean_text(response.choices[0].message.content)


def run_agent(prompt, content: str) -> str:
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": prompt},
            {"role": "user", "content": content},
        ],
        temperature=0,
    )
    return response.choices[0].message.content


def run_full_revision_via_agents(
    worksheet_text: str,
    markscheme_text: str,
    spec_text: str,
    on_step=None,
):
    combined_ws_ms = f"WORKSHEET:\n{worksheet_text}\n\nMARK SCHEME:\n{markscheme_text}"

    agent_steps = [
        ("Checking command word alignment...",   AGENT_1_PROMPT, f"WORKSHEET AND MARK SCHEME:\n{combined_ws_ms}"),
        ("Verifying mark allocations...",        AGENT_2_PROMPT, f"WORKSHEET AND MARK SCHEME:\n{combined_ws_ms}"),
        ("Checking physics accuracy...",         AGENT_3_PROMPT, f"WORKSHEET AND MARK SCHEME:\n{combined_ws_ms}"),
    ]
    coverage_input = f"WORKSHEET AND MARK SCHEME:\n{combined_ws_ms}\n\nINTENDED SCOPE:\n{spec_text}"
    agent_steps.append(("Evaluating topic coverage...", AGENT_4_PROMPT, coverage_input))

    reports = []
    total = 5
    for i, (label, prompt, content) in enumerate(agent_steps):
        if on_step:
            on_step(i, total, f"Agent {i+1}: {label}")
        reports.append(run_agent(prompt, content))

    report1, report2, report3, report4 = reports

    combined_input = f"""ORIGINAL WORKSHEET:
{worksheet_text}

ORIGINAL MARK SCHEME:
{markscheme_text}

INTENDED SCOPE:
{spec_text}

AGENT 1 REPORT:
{report1}

AGENT 2 REPORT:
{report2}

AGENT 3 REPORT:
{report3}

AGENT 4 REPORT:
{report4}
"""
    if on_step:
        on_step(4, total, "Agent 5: Intelligent revision and finalising...")
    return run_agent(AGENT_5_PROMPT, combined_input)


def parse_revised_output(text: str):
    ws_marker = "--- REVISED WORKSHEET ---"
    ms_marker = "--- REVISED MARK SCHEME ---"
    ws_idx = text.find(ws_marker)
    ms_idx = text.find(ms_marker)
    if ws_idx != -1 and ms_idx != -1:
        return text[ws_idx + len(ws_marker):ms_idx].strip(), text[ms_idx + len(ms_marker):].strip()
    ws_match = re.search(r"-{2,}\s*REVISED\s+WORKSHEET\s*-{2,}", text, re.IGNORECASE)
    ms_match = re.search(r"-{2,}\s*REVISED\s+MARK\s+SCHEME\s*-{2,}", text, re.IGNORECASE)
    if ws_match and ms_match:
        return text[ws_match.end():ms_match.start()].strip(), text[ms_match.end():].strip()
    return None, None


def run_formatting_agent(worksheet_text):
    cleaned = strip_answer_lines(clean_text(worksheet_text))
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": FORMATTING_AGENT_PROMPT},
            {"role": "user", "content": cleaned},
        ],
        temperature=0,
    )
    raw = response.choices[0].message.content.strip()
    raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.IGNORECASE)
    raw = re.sub(r"\s*```$", "", raw)
    try:
        return json.loads(raw)
    except json.JSONDecodeError as exc:
        raise ValueError(
            f"FormattingAgent returned invalid JSON ({exc}). "
            f"First 500 chars:\n{raw[:500]}"
        ) from exc


def render_formatted_preview(spec):
    """
    Render an HTML preview that matches the AQA-style GCSE exam paper format:
    - Question text on its own line (no inline marks)
    - Answer lines as underscore characters below the question
    - Marks on a separate right-aligned bold line after the answer lines
    - (Total for question X is Y marks) right-aligned bold
    """
    lines = spec.get("lines", [])
    st.markdown("#### Formatted Worksheet Preview")
    st.markdown("""
<style>
.worksheet-preview {
    max-width: 720px; padding: 20px 32px 20px 24px;
    border: 1px solid #2a2f3e; border-radius: 8px;
    background-color: #ffffff; color: #111;
    font-family: Arial, sans-serif; font-size: 11pt; line-height: 1.4;
}
.q-stem   { margin-top: 20px; padding-left: 0; }
.q-part   { margin-top: 10px; padding-left: 22px; }
.q-roman  { margin-top: 8px;  padding-left: 44px; }
.q-text   { color: #111; }
.q-ans-line {
    border-bottom: 1px solid #555;
    height: 18px; margin: 3px 0;
}
.q-ans-0  { margin-left: 0;    margin-right: 0; }
.q-ans-1  { margin-left: 11px; margin-right: 0; }
.q-ans-2  { margin-left: 22px; margin-right: 0; }
.q-marks  { text-align: right; font-weight: bold; color: #111;
            padding-left: 22px; margin-top: 2px; }
.q-total  { text-align: right; font-weight: bold; color: #111;
            margin-top: 8px; margin-bottom: 16px; }
/* Continuation lines for context+command-word splits — aligned to text start */
.q-cont-0 { padding-left: 24px;  margin-top: 1px; }
.q-cont-1 { padding-left: 46px;  margin-top: 1px; }
.q-cont-2 { padding-left: 68px;  margin-top: 1px; }
</style>""", unsafe_allow_html=True)

    # Regex to auto-split a context sentence from a command word.
    # Matches a command word that immediately follows sentence-ending punctuation.
    _CMD_SPLIT_RE = re.compile(
        r'(?<=[.?!])\s+(Calculate|Explain|State|Describe|Determine|Show|Find|'
        r'Compare|Evaluate|Suggest|Identify|Give|Write|Draw|Plot|Predict|Justify|'
        r'Define|Outline|Use|Work\s+out|Complete|Name|Tick|Circle|Underline|Label)\b',
        re.IGNORECASE
    )

    html_lines = ['<div class="worksheet-preview">']
    last_q = None
    _COMB_RE = re.compile(r'^\(([a-z])\) \(([ivxlcdm]+)\)$', re.IGNORECASE)
    _seen_roman_parents: dict = {}  # {qnum: set_of_parent_letters_already_shown}
    _orphan_qnum = None  # level-0 number with no text — merged onto next line

    for line in lines:
        qnum = line.get("question_number")
        indent_level = int(line.get("indent_level", 0))
        part_label = line.get("part_label") or ""
        subpart_label = line.get("subpart_label") or ""
        question_text = line.get("question_text") or ""
        if question_text.strip().lower() == "none":
            question_text = ""
        marks = line.get("marks")
        is_total = bool(line.get("is_total_for_question"))

        if is_total:
            html_lines.append(
                f'<div class="q-total">'
                f'(Total for question {qnum} is {marks} marks)</div>'
            )
            last_q = qnum
            _orphan_qnum = None
            continue

        # Level-0 items with no question text are just question numbers (e.g. "1").
        # Save the number and merge it onto the next line so "1  (a)(i)  text"
        # all appear together instead of "1" alone on its own line.
        if indent_level == 0 and not question_text.strip():
            _orphan_qnum = qnum
            last_q = qnum
            continue

        # ── Roman numeral de-duplication ────────────────────────────────────
        _eff_level = indent_level
        _eff_label = part_label
        if indent_level == 1:
            _mc = _COMB_RE.match(part_label)
            if _mc:
                _parent = _mc.group(1)
                _roman  = _mc.group(2)
                _seen_roman_parents.setdefault(qnum, set())
                if _parent in _seen_roman_parents[qnum]:
                    _eff_level = 2
                    _eff_label = f"({_roman})"
                else:
                    _seen_roman_parents[qnum].add(_parent)

        # ── Merge orphan question number ─────────────────────────────────────
        # If the previous level-0 had no text, prepend its number to this label.
        _ans_level = _eff_level  # used for answer-line indentation
        if _orphan_qnum is not None:
            _eff_label = (f"{_orphan_qnum}  {_eff_label}" if _eff_label
                          else str(_orphan_qnum))
            _eff_level = 0   # render the merged line at stem indentation
            _orphan_qnum = None

        # CSS class for indent level
        if _eff_level == 0:
            div_class = "q-stem"
        elif _eff_level == 1:
            div_class = "q-part"
        else:
            div_class = "q-roman"

        # Build label (plain text — no bold per AQA exam style)
        if _eff_level == 0:
            lbl = _eff_label or (str(qnum) if qnum else "")
            label_html = f'{lbl}&nbsp;&nbsp;' if lbl else ""
        elif _eff_level == 1 and _eff_label:
            label_html = f'{_eff_label}&nbsp;'
        elif _eff_level >= 2 and (subpart_label or _eff_label):
            label_html = f'{subpart_label or _eff_label}&nbsp;'
        else:
            label_html = ""

        # Auto-split context sentence from command word, then handle any \n in text
        question_text = _CMD_SPLIT_RE.sub(
            lambda m: '\n' + m.group(0).lstrip(), question_text
        )
        _qt_parts = [s for s in question_text.split('\n') if s.strip()]

        # First part: rendered with label
        _first = _qt_parts[0] if _qt_parts else ""
        html_lines.append(
            f'<div class="{div_class} q-text">{label_html}{_first}</div>'
        )
        # Continuation parts (command words) aligned to text-start position
        _cont_cls = f"q-cont-{min(_ans_level, 2)}"
        for _qtp in _qt_parts[1:]:
            html_lines.append(f'<div class="{_cont_cls} q-text">{_qtp}</div>')

        # Answer lines + marks below
        # Line count: 1 mark=1 line, 2=2, 3=3, 4=4, 5=5, 6+=6 (max 6)
        if marks and marks > 0:
            m = int(marks)
            num_lines = min(m, 6)
            ans_class = f"q-ans-{min(_ans_level, 2)}"
            for _ in range(num_lines):
                html_lines.append(f'<div class="q-ans-line {ans_class}"></div>')
            # Marks on separate right-aligned line
            html_lines.append(f'<div class="q-marks">({marks})</div>')

        last_q = qnum

    html_lines.append("</div>")
    st.markdown("\n".join(html_lines), unsafe_allow_html=True)


# ================================================================
# DOCX HELPERS
# ================================================================

def _set_run_font(run, bold=False, size_pt=11):
    run.font.name = "Arial"
    run.font.size = Pt(size_pt)
    run.bold = bold
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Arial')
    rFonts.set(qn('w:hAnsi'), 'Arial')
    rPr.insert(0, rFonts)


def build_formatted_docx(spec):
    """
    Build a fully formatted A4 Word document matching AQA GCSE exam paper style.

    Layout (matches reference):
      • Question text on its own line with label (no inline marks)
      • Answer lines as underscore-character paragraphs below the question
      • Marks (n) on a separate RIGHT-ALIGNED BOLD paragraph after answer lines
      • (Total for question X is Y marks) — right-aligned bold
      • No paragraph borders on question text paragraphs
    """
    document = Document()

    # --- Page setup (A4, 1.91cm side margins matching reference 1080 twips) ---
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width  = Cm(21.0)
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(1.91)   # 1080 twips
    section.right_margin  = Cm(1.91)   # 1080 twips

    # --- Normal style ---
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.paragraph_format.space_after  = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # --- Indentation constants (cm, from reference XML) ---
    # Level 0 (main question):  label starts at 0, text wraps at 0.7cm
    # Level 1 ((a)(b)(c)):      label at 0.63cm (718-359 twips), text at 1.27cm
    # Level 2 ((i)(ii)(iii)):   label at 1.27cm (1078-358 twips), text at 1.90cm
    LABEL_CM = [0.0,  0.63, 1.27]
    TEXT_CM  = [0.7,  1.27, 1.90]

    # Answer line: underscore string filling ~17cm content width.
    # Lines start at LABEL_CM (not TEXT_CM) so they begin under the label letter/roman.
    # Counts adjusted for the wider start position (LABEL_CM is further left than TEXT_CM).
    ANSWER_UNDERSCORES = {0: 80, 1: 77, 2: 73}

    # Auto-split context sentence from command word inside question_text
    _CMD_SPLIT_D = re.compile(
        r'(?<=[.?!])\s+(Calculate|Explain|State|Describe|Determine|Show|Find|'
        r'Compare|Evaluate|Suggest|Identify|Give|Write|Draw|Plot|Predict|Justify|'
        r'Define|Outline|Use|Work\s+out|Complete|Name|Tick|Circle|Underline|Label)\b',
        re.IGNORECASE
    )

    lines         = spec.get("lines", [])
    paper_total   = spec.get("paper_total_marks")
    last_q        = None
    prev_indent   = None
    _COMB_RE_D    = re.compile(r'^\(([a-z])\) \(([ivxlcdm]+)\)$', re.IGNORECASE)
    _seen_roman_d: dict = {}  # {qnum: set_of_parent_letters_already_shown}
    _orphan_qnum_d = None  # level-0 number with no text — merged onto next line

    def _para(space_before_pt=0, space_after_pt=0,
               left_cm=0.0, hanging_cm=0.0,
               align=WD_ALIGN_PARAGRAPH.LEFT):
        p = document.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(space_before_pt)
        pf.space_after  = Pt(space_after_pt)
        pf.alignment    = align
        if left_cm:
            pf.left_indent = Cm(left_cm)
        if hanging_cm:
            pf.first_line_indent = Cm(-hanging_cm)
        return p

    for line in lines:
        qnum        = line.get("question_number")
        indent_level = int(line.get("indent_level", 0))
        part_label   = line.get("part_label")   or ""
        subpart_label= line.get("subpart_label") or ""
        question_text= line.get("question_text") or ""
        if question_text.strip().lower() == "none":
            question_text = ""
        marks    = line.get("marks")
        is_total = bool(line.get("is_total_for_question"))

        # ── (Total for question X) ──────────────────────────────────────────
        if is_total:
            p = _para(space_before_pt=4, space_after_pt=12,
                      align=WD_ALIGN_PARAGRAPH.RIGHT)
            r = p.add_run(f"(Total for question {qnum} is {marks} marks)")
            _set_run_font(r, bold=True)
            last_q = qnum
            prev_indent = None
            _orphan_qnum_d = None
            continue

        # ── Level-0 items with no text — save number to merge onto next line ─
        if indent_level == 0 and not question_text.strip():
            _orphan_qnum_d = qnum
            last_q = qnum
            continue

        # ── Roman numeral de-duplication ────────────────────────────────────
        _eff_indent = indent_level
        _eff_part   = part_label
        if indent_level == 1:
            _mcd = _COMB_RE_D.match(part_label)
            if _mcd:
                _parent_d = _mcd.group(1)
                _roman_d  = _mcd.group(2)
                _seen_roman_d.setdefault(qnum, set())
                if _parent_d in _seen_roman_d[qnum]:
                    _eff_indent = 2
                    _eff_part   = f"({_roman_d})"
                else:
                    _seen_roman_d[qnum].add(_parent_d)

        # ── Merge orphan question number ─────────────────────────────────────
        _ans_lvl_d = _eff_indent  # separate tracker for answer-line indent
        if _orphan_qnum_d is not None:
            _eff_part   = (f"{_orphan_qnum_d}  {_eff_part}" if _eff_part
                           else str(_orphan_qnum_d))
            _eff_indent = 0   # render merged line at stem indentation
            _orphan_qnum_d = None

        # ── Spacing before this paragraph ───────────────────────────────────
        if _eff_indent == 0:
            sp = 22 if (last_q is not None and qnum != last_q) else 0
        elif _eff_indent == 1:
            sp = 10 if (prev_indent is not None and prev_indent != 0) else 6
        else:  # level 2
            sp = 8

        # ── Label (plain text — no bold per AQA exam style) ─────────────────
        if _eff_indent == 0:
            label      = _eff_part or (str(qnum) if qnum else "")
        elif _eff_indent == 1 and _eff_part:
            label = _eff_part
        elif _eff_indent >= 2 and (subpart_label or _eff_part):
            label = subpart_label or _eff_part
        else:
            label = ""
        label_bold = False  # AQA style: labels are plain, not bold

        # ── Question text paragraph ─────────────────────────────────────────
        lvl = min(_eff_indent, 2)
        # For merged orphan items (_eff_indent forced to 0 but child level tracked in
        # _ans_lvl_d), use the child level's TEXT_CM as the wrap-back point so
        # continuation lines align under the first word, not under the "1" label.
        _child_lvl = min(_ans_lvl_d, 2)
        if lvl == 0 and _child_lvl > 0:
            # Was a merged orphan — wrap at child TEXT_CM
            _left_cm = TEXT_CM[_child_lvl]
            _hang_cm = TEXT_CM[_child_lvl]  # first line at 0, wrap at TEXT_CM[child]
        else:
            _left_cm = TEXT_CM[lvl]
            _hang_cm = TEXT_CM[lvl] - LABEL_CM[lvl]
        # Auto-split context sentence from command word, then handle \n splits
        question_text = _CMD_SPLIT_D.sub(
            lambda m: '\n' + m.group(0).lstrip(), question_text
        )
        _qt_parts = [s for s in question_text.split('\n') if s.strip()]

        # First paragraph — label + first part of question text
        qp = _para(space_before_pt=sp, space_after_pt=0,
                   left_cm=_left_cm, hanging_cm=_hang_cm)
        if label:
            r_lbl = qp.add_run(label + "  ")
            _set_run_font(r_lbl, bold=label_bold)
        if _qt_parts:
            r_txt = qp.add_run(_qt_parts[0])
            _set_run_font(r_txt, bold=False)

        # Continuation paragraphs (command word lines) — no hanging indent,
        # start at TEXT_CM[lvl] so they align with the first word above.
        for _qtp in _qt_parts[1:]:
            cp = _para(space_before_pt=0, space_after_pt=0,
                       left_cm=TEXT_CM[min(_eff_indent, 2)])
            r_cp = cp.add_run(_qtp)
            _set_run_font(r_cp, bold=False)

        # ── Answer lines (underscore text) ──────────────────────────────────
        if marks and marks > 0:
            m = int(marks)
            num_ans = min(m, 6)  # 1 mark=1 line, 2=2, ..., 6+=6
            # Use _ans_lvl_d for underscore count/indent so merged lines still
            # get the correct answer-line width (not the collapsed lvl=0 width).
            _alvl = min(_ans_lvl_d, 2)
            underscores = "_" * ANSWER_UNDERSCORES.get(_alvl, 72)

            for i in range(num_ans):
                ap = _para(space_before_pt=0, space_after_pt=0,
                           left_cm=LABEL_CM[_alvl])
                r = ap.add_run(underscores)
                _set_run_font(r, bold=False)

            # ── Marks on their own right-aligned line ────────────────────────
            mp = _para(space_before_pt=0, space_after_pt=2,
                       left_cm=LABEL_CM[_alvl],
                       align=WD_ALIGN_PARAGRAPH.RIGHT)
            r_m = mp.add_run(f"({marks})")
            _set_run_font(r_m, bold=True)

        last_q      = qnum
        prev_indent = _eff_indent

    # ── Paper total ─────────────────────────────────────────────────────────
    if paper_total:
        p = _para(space_before_pt=14, align=WD_ALIGN_PARAGRAPH.RIGHT)
        r = p.add_run(f"Total marks for question paper: {paper_total}")
        _set_run_font(r, bold=True)
        r.underline = True

    bio = BytesIO()
    document.save(bio)
    bio.seek(0)
    return bio


def build_markscheme_docx(markscheme_text: str) -> BytesIO:
    """
    Build a mark scheme DOCX matching the Edexcel format:
    - Each marking point ends with bold (1)
    - 'Any X from:' lines followed by bullet points
    - (Total for question X is Y marks) — bold, 'is' not '='
    - Total marks line — bold + underlined
    """
    document = Document()
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(1.91)
    section.right_margin = Cm(1.91)
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    MAIN_Q_RE = re.compile(r"^\s*(\d+)\s")
    PART_RE = re.compile(r"^\s*\(([a-z])\)")
    ROMAN_RE = re.compile(r"^\s*\((i{1,4}|iv|vi{0,3}|ix|xi{0,3}|x{1,3})\)", re.IGNORECASE)
    BULLET_RE = re.compile(r"^\s*[•\-\*]\s+")
    TOTAL_Q_RE = re.compile(r"\(Total for question", re.IGNORECASE)
    TOTAL_PAPER_RE = re.compile(r"Total marks for question paper", re.IGNORECASE)
    TOTAL_PAPER_ALT = re.compile(r"Total for paper\s*[=:]", re.IGNORECASE)

    def add_inline_bold_marks(p, text, base_bold=False):
        """Split text on (1) and emit alternating normal / bold runs."""
        parts = re.split(r'(\(1\))', text)
        for part in parts:
            if part == "(1)":
                r = p.add_run("(1)")
                _set_run_font(r, bold=True)
            elif part:
                r = p.add_run(part)
                _set_run_font(r, bold=base_bold)

    prev_line_type = None  # 'main', 'part', 'bullet', 'other', 'total'

    for raw_line in markscheme_text.split("\n"):
        line = raw_line.strip()
        if not line:
            continue

        # ---- Total for paper ----
        if TOTAL_PAPER_RE.search(line) or TOTAL_PAPER_ALT.search(line):
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(line)
            _set_run_font(r, bold=True)
            r.underline = True
            prev_line_type = 'total'
            continue

        # ---- Total for question ----
        if TOTAL_Q_RE.search(line):
            # Normalise "=" -> "is"
            line = re.sub(
                r'(Total for question\s+\w+)\s*=\s*(\d+)',
                r'\1 is \2',
                line, flags=re.IGNORECASE
            )
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(12)
            r = p.add_run(line)
            _set_run_font(r, bold=True)
            prev_line_type = 'total'
            continue

        # ---- Bullet point ----
        if BULLET_RE.match(line):
            text_after_bullet = BULLET_RE.sub("", line).strip()
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Cm(1.8)
            p.paragraph_format.first_line_indent = Cm(-0.4)
            r = p.add_run("\u2022 ")
            _set_run_font(r, bold=False)
            add_inline_bold_marks(p, text_after_bullet)
            prev_line_type = 'bullet'
            continue

        # ---- Identify line type ----
        is_main = MAIN_Q_RE.match(line)
        is_part = PART_RE.match(line) or ROMAN_RE.match(line)

        if is_main:
            sp_before = 14 if prev_line_type in ('total', 'part', 'bullet', 'other') else 0
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(sp_before)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Cm(0)
            m = is_main
            r = p.add_run(m.group(1) + " ")
            _set_run_font(r, bold=True)
            add_inline_bold_marks(p, line[m.end():])
            prev_line_type = 'main'

        elif is_part:
            sp_before = 4 if prev_line_type in ('part', 'bullet', 'other') else 0
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(sp_before)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Cm(0.8)
            add_inline_bold_marks(p, line)
            prev_line_type = 'part'

        else:
            # Continuation / any-from line / notes
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Cm(1.5)
            add_inline_bold_marks(p, line)
            prev_line_type = 'other'

    bio = BytesIO()
    document.save(bio)
    bio.seek(0)
    return bio


# ================================================================
# MAIN PIPELINE
# ================================================================

if run_button and worksheet_file:
    _PIPELINE_STEPS = [
        ("📂", "Read Files",      "Parsing documents"),
        ("✏️",  "Enhance WS",     "Improving worksheet"),
        ("📋", "Mark Scheme",    "Generating MS"),
        ("🔍", "Agent 1",        "Command words"),
        ("📐", "Agent 2",        "Mark structure"),
        ("🔬", "Agent 3",        "Cognitive balance"),
        ("🗂️", "Agent 4",        "Topic coverage"),
        ("✨", "Agent 5",        "Final revision"),
    ]
    _N = len(_PIPELINE_STEPS)

    def _render_pipeline(current_step, step_label=""):
        pct = int((current_step / _N) * 100)
        cards_html = ""
        for i, (icon, short, detail) in enumerate(_PIPELINE_STEPS):
            if i < current_step:
                state = "done"
                show_icon = "✓"
                check_html = f'<div class="card-check">✓</div>'
            elif i == current_step:
                state = "active"
                show_icon = icon
                check_html = ""
            else:
                state = "pending"
                show_icon = icon
                check_html = ""
            cards_html += (
                f'<div class="agent-card {state}">'
                f'{check_html}'
                f'<div class="card-icon">{show_icon}</div>'
                f'<div class="card-label">{short}</div>'
                f'<div class="card-status-dot"></div>'
                f'</div>'
            )
        step_desc = step_label or (f"Step {current_step+1} of {_N}" if current_step < _N else "Complete")
        return (
            f'<div class="pipeline-wrap">'
            f'<div class="pipeline-title">⚙️ Enhancement Pipeline'
            f'<span style="font-size:0.75rem;font-weight:400;color:{_text_muted};margin-left:8px">'
            f'{step_desc}</span></div>'
            f'<div class="cards-grid">{cards_html}</div>'
            f'<div class="pipeline-progress">'
            f'<div class="pipeline-progress-bar" style="width:{pct}%"></div>'
            f'</div>'
            f'</div>'
        )

    _prog_box = st.empty()

    # Step 0: read files
    _prog_box.markdown(_render_pipeline(0, "Parsing uploaded documents…"), unsafe_allow_html=True)
    worksheet_text = extract_docx(worksheet_file)
    markscheme_text = extract_docx(markscheme_file) if markscheme_file else ""
    spec_text = read_spec_text(spec_txt, spec_docx, pasted_spec)

    # Step 1: enhance worksheet
    _prog_box.markdown(_render_pipeline(1, "Enhancing worksheet quality…"), unsafe_allow_html=True)
    improved_ws = improve_worksheet(worksheet_text)

    # Step 2: generate mark scheme
    _prog_box.markdown(_render_pipeline(2, "Generating mark scheme…"), unsafe_allow_html=True)
    improved_ms = generate_markscheme(improved_ws)

    # Step 3–6: Agents 1–4
    _combined = f"WORKSHEET:\n{improved_ws}\n\nMARK SCHEME:\n{improved_ms}"
    _prog_box.markdown(_render_pipeline(3, "Agent 1: Checking command word alignment…"), unsafe_allow_html=True)
    _r1 = run_agent(AGENT_1_PROMPT, f"WORKSHEET AND MARK SCHEME:\n{_combined}")

    _prog_box.markdown(_render_pipeline(4, "Agent 2: Verifying mark scheme structure…"), unsafe_allow_html=True)
    _r2 = run_agent(AGENT_2_PROMPT, f"WORKSHEET AND MARK SCHEME:\n{_combined}")

    _prog_box.markdown(_render_pipeline(5, "Agent 3: Evaluating cognitive balance…"), unsafe_allow_html=True)
    _r3 = run_agent(AGENT_3_PROMPT, f"WORKSHEET AND MARK SCHEME:\n{_combined}")

    _prog_box.markdown(_render_pipeline(6, "Agent 4: Checking topic coverage…"), unsafe_allow_html=True)
    _r4 = run_agent(AGENT_4_PROMPT,
        f"WORKSHEET AND MARK SCHEME:\n{_combined}\n\nINTENDED SCOPE:\n{spec_text}")

    # Step 7: Agent 5 — intelligent revision
    _prog_box.markdown(_render_pipeline(7, "Agent 5: Applying intelligent revision…"), unsafe_allow_html=True)
    _agent5_input = (
        f"ORIGINAL WORKSHEET:\n{improved_ws}\n\nORIGINAL MARK SCHEME:\n{improved_ms}\n\n"
        f"INTENDED SCOPE:\n{spec_text}\n\n"
        f"AGENT 1 REPORT:\n{_r1}\n\nAGENT 2 REPORT:\n{_r2}\n\n"
        f"AGENT 3 REPORT:\n{_r3}\n\nAGENT 4 REPORT:\n{_r4}"
    )
    _final_text = run_agent(AGENT_5_PROMPT, _agent5_input)
    _revised_ws, _revised_ms = parse_revised_output(_final_text)
    if _revised_ws:
        improved_ws = _revised_ws
    if _revised_ms:
        improved_ms = _revised_ms

    # Show all steps complete
    _prog_box.markdown(
        _render_pipeline(_N, "All steps complete!") +
        f'<div style="background:{_success};border-radius:10px;padding:14px 22px;'
        f'margin-top:12px;font-family:Inter,Arial,sans-serif;font-size:14px;'
        f'color:#ffffff;font-weight:600;display:flex;align-items:center;gap:10px;">'
        f'<span style="font-size:1.3rem">✅</span>'
        f'Enhancement complete — review your results and export below.'
        f'</div>',
        unsafe_allow_html=True,
    )

    # Derive default download filenames from the uploaded worksheet filename.
    # Replace "QP" with "MS" (case-insensitive) for the mark scheme name.
    _ws_fname_default = getattr(worksheet_file, "name", "gcse_worksheet_formatted.docx")
    _ms_fname_default = re.sub(r'\bQP\b', 'MS', _ws_fname_default, flags=re.IGNORECASE)
    if _ms_fname_default == _ws_fname_default:
        # No "QP" found — append _MS before .docx
        _ms_fname_default = re.sub(r'(\.docx)$', '_MS\\1', _ws_fname_default,
                                   flags=re.IGNORECASE) or _ws_fname_default.replace('.docx', '_MS.docx')

    st.session_state.update({
        "worksheet_text": worksheet_text,
        "markscheme_text": markscheme_text,
        "improved_ws": improved_ws,
        "improved_ms": improved_ms,
        "spec_text": spec_text,
        "ws_fname": _ws_fname_default,
        "ms_fname": _ms_fname_default,
    })
    for k in ("fmt_spec", "fmt_docx_bytes", "ms_docx_bytes"):
        st.session_state.pop(k, None)


# ================================================================
# OUTPUT SECTION
# ================================================================

if "worksheet_text" in st.session_state and st.session_state["worksheet_text"]:
    worksheet_text = st.session_state["worksheet_text"]
    markscheme_text = st.session_state.get("markscheme_text", "")
    improved_ws = st.session_state.get("improved_ws", "")
    improved_ms = st.session_state.get("improved_ms", "")
    spec_text = st.session_state.get("spec_text", "")

    # Apply any AI-pending updates to widget keys BEFORE the text area widgets
    # render — Streamlit forbids setting a widget key after that widget renders.
    if "_ws_pending" in st.session_state:
        st.session_state["ws_editor"] = st.session_state.pop("_ws_pending")
    if "_ms_pending" in st.session_state:
        st.session_state["ms_editor"] = st.session_state.pop("_ms_pending")

    st.markdown(f"""
<div class="section-row">
  <div class="section-bar" style="background:{_primary};"></div>
  <div>
    <div class="section-title">Enhanced Worksheet</div>
    <div class="section-sub">Edit directly below — click Save Edits to apply changes</div>
  </div>
</div>
""", unsafe_allow_html=True)
    _ws_edit = st.text_area("Worksheet Output", value=st.session_state.get("improved_ws", improved_ws),
                            height=400, key="ws_editor", label_visibility="collapsed")
    if st.button("💾 Save Worksheet Edits", key="save_ws_edit"):
        st.session_state["improved_ws"] = _ws_edit
        for k in ("fmt_spec", "fmt_docx_bytes"):
            st.session_state.pop(k, None)
        st.success("✅ Worksheet edits saved.")

    st.markdown(f"""
<div class="section-row">
  <div class="section-bar" style="background:{_violet};"></div>
  <div>
    <div class="section-title">Enhanced Mark Scheme</div>
    <div class="section-sub">Edit directly below — click Save Edits to apply changes</div>
  </div>
</div>
""", unsafe_allow_html=True)
    _ms_edit = st.text_area("Mark Scheme Output", value=st.session_state.get("improved_ms", improved_ms),
                            height=400, key="ms_editor", label_visibility="collapsed")
    if st.button("💾 Save Mark Scheme Edits", key="save_ms_edit"):
        st.session_state["improved_ms"] = _ms_edit
        st.session_state.pop("ms_docx_bytes", None)
        st.success("✅ Mark scheme edits saved.")

    st.markdown(f'<div style="border-top:1px solid {_border};margin:28px 0 18px 0;"></div>', unsafe_allow_html=True)

    # ---- QA Validation ----
    with st.expander("🔎 QA Validation Report", expanded=False):
        misaligned = False
        validation_ms_text = improved_ms or markscheme_text

        if validation_ms_text:
            overlap = keyword_overlap(improved_ws, validation_ms_text)
            st.markdown(f"""
<div class="qa-card" style="display:flex;align-items:center;gap:16px;">
  <div style="font-size:2rem;">{'🟢' if overlap >= 40 else '🔴'}</div>
  <div>
    <div style="font-size:0.8rem;font-weight:700;color:{_text_muted};">KEYWORD ALIGNMENT</div>
    <div style="font-size:1.4rem;font-weight:800;color:{_success if overlap >= 40 else _error};">{overlap}%</div>
  </div>
</div>
""", unsafe_allow_html=True)
            if overlap < 40:
                st.error("⚠️ Content misalignment detected — low keyword overlap between worksheet and mark scheme.")

        ws_total = extract_total(improved_ws)
        ms_total = extract_total(validation_ms_text)
        if ws_total and ms_total and ws_total != ms_total:
            st.error(f"❌ Total mark mismatch: Worksheet = {ws_total}, Mark Scheme = {ms_total}")
            misaligned = True
        if fractional_marks_present(validation_ms_text):
            st.error("❌ Fractional marks detected in mark scheme.")
            misaligned = True

        ws_questions = extract_question_numbers(improved_ws)
        ms_questions = extract_question_numbers(validation_ms_text)
        mismatch_details = []
        if ws_questions != ms_questions:
            missing = [q for q in ws_questions if q not in ms_questions]
            extra   = [q for q in ms_questions if q not in ws_questions]
            if missing:
                mismatch_details.append(f"Missing from mark scheme: {missing}")
            if extra:
                mismatch_details.append(f"Extra in mark scheme (remove): {extra}")
            st.error("❌ Question number mismatch detected.")
            st.write(f"Worksheet Qs: {ws_questions}  |  Mark Scheme Qs: {ms_questions}")
            misaligned = True
        else:
            st.success("✅ Question numbers align correctly.")

        mismatch_info_str = "\n".join(mismatch_details) if mismatch_details else None

        if misaligned:
            if st.button("🔄 Regenerate Mark Scheme from Worksheet"):
                regenerated = generate_markscheme(improved_ws, mismatch_info=mismatch_info_str)
                st.session_state["improved_ms"] = regenerated
                st.session_state.pop("ms_docx_bytes", None)
                st.text_area("Regenerated Mark Scheme", regenerated, height=400)
                st.success("✅ Mark scheme regenerated and saved.")
        else:
            st.success("✅ All structural checks passed.")

    st.markdown(f'<div style="border-top:1px solid {_border};margin:18px 0 18px 0;"></div>', unsafe_allow_html=True)

    # ================================================================
    # EXPORT
    # ================================================================
    st.markdown(f"""
<div class="section-row">
  <div class="section-bar" style="background:{_success};"></div>
  <div>
    <div class="section-title">Export Documents</div>
    <div class="section-sub">Generate formatted .docx files ready to share with students</div>
  </div>
</div>
""", unsafe_allow_html=True)

    validation_ms_text = improved_ms or markscheme_text
    misaligned_for_export = False
    if validation_ms_text:
        ws_total = extract_total(improved_ws)
        ms_total = extract_total(validation_ms_text)
        if ws_total and ms_total and ws_total != ms_total:
            misaligned_for_export = True
        if fractional_marks_present(validation_ms_text):
            misaligned_for_export = True
        if extract_question_numbers(improved_ws) != extract_question_numbers(validation_ms_text):
            misaligned_for_export = True

    override_ok = True
    if misaligned_for_export:
        st.warning("QA checks found structural issues. You can still export — double-check the output.")
        override_ok = st.checkbox("Proceed with export despite QA warnings", key="fmt_override")

    if override_ok:
        # ── Filename inputs ──────────────────────────────────────────────────
        st.markdown(f'<div class="section-label" style="margin-bottom:8px;">Output Filenames</div>', unsafe_allow_html=True)
        _fcol1, _fcol2 = st.columns(2)
        with _fcol1:
            _ws_fname = st.text_input(
                "Worksheet filename",
                value=st.session_state.get("ws_fname", "gcse_worksheet_formatted.docx"),
                key="ws_fname_input",
            )
        with _fcol2:
            _ms_fname = st.text_input(
                "Mark scheme filename",
                value=st.session_state.get("ms_fname", "gcse_markscheme_formatted.docx"),
                key="ms_fname_input",
            )

        _exp_col1, _exp_col2 = st.columns(2)

        with _exp_col1:
            st.markdown(f"""
<div class="export-card">
  <div class="export-card-stripe" style="background:{_primary};"></div>
  <div class="export-card-body">
    <div class="export-card-title">📄  Formatted Worksheet</div>
    <div class="export-card-desc">Exam-paper style layout with answer lines, marks placement and question numbering.</div>
  </div>
</div>
""", unsafe_allow_html=True)
            if st.button("⚡ Generate Formatted Worksheet", key="fmt_ws", use_container_width=True):
                with st.spinner("Running FormattingAgent — structuring layout..."):
                    try:
                        fmt_spec = run_formatting_agent(improved_ws)
                        docx_bytes = build_formatted_docx(fmt_spec)
                        st.session_state["fmt_spec"] = fmt_spec
                        st.session_state["fmt_docx_bytes"] = docx_bytes
                    except Exception as e:
                        st.error(f"Worksheet export failed: {e}")

            if "fmt_docx_bytes" in st.session_state:
                st.download_button(
                    label="⬇  Download Worksheet (.docx)",
                    data=st.session_state["fmt_docx_bytes"],
                    file_name=_ws_fname or "gcse_worksheet_formatted.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl_ws",
                    use_container_width=True,
                )

        with _exp_col2:
            st.markdown(f"""
<div class="export-card">
  <div class="export-card-stripe" style="background:{_success};"></div>
  <div class="export-card-body">
    <div class="export-card-title">📋  Mark Scheme</div>
    <div class="export-card-desc">Clean Word document with bold question numbers and structured marking points.</div>
  </div>
</div>
""", unsafe_allow_html=True)
            if st.button("⚡ Generate Mark Scheme (.docx)", key="fmt_ms", use_container_width=True):
                with st.spinner("Building mark scheme document..."):
                    try:
                        ms_bytes = build_markscheme_docx(st.session_state.get("improved_ms", improved_ms))
                        st.session_state["ms_docx_bytes"] = ms_bytes
                    except Exception as e:
                        st.error(f"Mark scheme export failed: {e}")

            if "ms_docx_bytes" in st.session_state:
                st.download_button(
                    label="⬇  Download Mark Scheme (.docx)",
                    data=st.session_state["ms_docx_bytes"],
                    file_name=_ms_fname or "gcse_markscheme_formatted.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl_ms",
                    use_container_width=True,
                )

        if "fmt_spec" in st.session_state:
            st.markdown(f'<div style="border-top:1px solid {_border};margin:20px 0 16px 0;"></div>', unsafe_allow_html=True)
            render_formatted_preview(st.session_state["fmt_spec"])


# ================================================================
# AI CHAT ASSISTANT
# ================================================================

st.markdown(f'<hr>', unsafe_allow_html=True)
st.markdown(f"""
<div class="section-row" style="margin-top:8px;">
  <div class="section-bar" style="background:{_violet};"></div>
  <div>
    <div class="section-title">💬 AI Assistant</div>
    <div class="section-sub">Make targeted edits — changes apply automatically to your documents</div>
  </div>
</div>
<div class="chat-prompt-chips">
  <span class="chat-chip">"Change Olivia to Mustafa throughout"</span>
  <span class="chat-chip">"Make Q3b an explain question, 4 marks"</span>
  <span class="chat-chip">"Fix the calculation in question 2"</span>
</div>
""", unsafe_allow_html=True)

_CHAT_SYSTEM = """You are an expert GCSE Physics worksheet editor.
You help teachers improve their worksheets and mark schemes.

IMPORTANT: You MUST respond with ONLY valid JSON — no prose, no markdown fences, nothing outside the JSON object.

JSON schema:
{
  "message": "<short explanation of what you did or why>",
  "action": "modify" | "info",
  "rerun_pipeline": true | false,
  "changes": [
    {
      "target": "worksheet" | "markscheme",
      "find": "<exact text to find, verbatim from the document>",
      "replace": "<exact replacement text>"
    }
  ]
}

Rules:
- Use action "modify" when making any edit to the worksheet or mark scheme.
- Use action "info" for questions, explanations or when no edit is needed (changes array will be empty).
- Set rerun_pipeline to FALSE for virtually everything: renaming a student/character, changing a command word (e.g. "state" → "explain"), fixing a number, rewording a question, correcting a mark allocation, or any targeted single edit.
- Set rerun_pipeline to TRUE ONLY when the user explicitly asks to REGENERATE or COMPLETELY RESTRUCTURE the whole worksheet from scratch — e.g. "change the topic to waves", "add five new questions", "rebuild the whole worksheet". Simple word/name/phrasing changes are NEVER a reason to rerun the pipeline.
- "find" must be an exact verbatim substring of the current document — copy it exactly.
- "replace" is the new text that replaces that exact substring.
- You may include multiple change objects in the changes array (e.g. one for worksheet, one for markscheme).
- If action is "info", set changes to [] and rerun_pipeline to false.
"""

if "chat_history" not in st.session_state:
    st.session_state["chat_history"] = []


def _apply_chat_changes(changes, ws, ms):
    """Apply a list of {target, find, replace} edits. Returns (new_ws, new_ms, applied_count)."""
    applied = 0
    for ch in changes:
        target = ch.get("target", "worksheet")
        find_str = ch.get("find", "")
        replace_str = ch.get("replace", "")
        if not find_str:
            continue
        if target == "worksheet" and find_str in ws:
            ws = ws.replace(find_str, replace_str, 1)
            applied += 1
        elif target == "markscheme" and find_str in ms:
            ms = ms.replace(find_str, replace_str, 1)
            applied += 1
    return ws, ms, applied


_chat_container = st.container()
with _chat_container:
    for msg in st.session_state["chat_history"]:
        display_text = msg.get("display", msg["content"])
        if msg["role"] == "user":
            st.markdown(
                f'<div class="chat-user">'
                f'<div class="chat-name" style="color:{_primary};">You</div>'
                f'{display_text}</div>',
                unsafe_allow_html=True,
            )
        else:
            st.markdown(
                f'<div class="chat-assistant">'
                f'<div class="chat-name" style="color:{_violet};">AI Assistant</div>'
                f'{display_text}</div>',
                unsafe_allow_html=True,
            )

with st.form("chat_form", clear_on_submit=True):
    _cols = st.columns([6, 1])
    _user_input = _cols[0].text_input(
        "Message", placeholder="Ask me to edit your worksheet or mark scheme…", label_visibility="collapsed"
    )
    _send = _cols[1].form_submit_button("Send ➤", use_container_width=True)

if _send and _user_input.strip():
    _ctx_ws = st.session_state.get("improved_ws", "")
    _ctx_ms = st.session_state.get("improved_ms", "")
    _context_block = ""
    if _ctx_ws:
        _context_block = (
            f"\n\nCurrent worksheet:\n{_ctx_ws[:4000]}"
            + (f"\n\nCurrent mark scheme:\n{_ctx_ms[:2500]}" if _ctx_ms else "")
        )

    _messages = [{"role": "system", "content": _CHAT_SYSTEM + _context_block}]
    for _m in st.session_state["chat_history"][-10:]:
        _messages.append({"role": _m["role"], "content": _m["content"]})
    _messages.append({"role": "user", "content": _user_input.strip()})

    st.session_state["chat_history"].append({"role": "user", "content": _user_input.strip(), "display": _user_input.strip()})

    with st.spinner("🤖 Assistant is thinking…"):
        _resp = client.chat.completions.create(
            model="gpt-4o",
            messages=_messages,
            max_tokens=1500,
            temperature=0.3,
        )
        _raw = _resp.choices[0].message.content.strip()

    # Parse JSON response
    try:
        # Strip markdown fences if model wrapped it anyway
        _clean = re.sub(r"^```(?:json)?\s*|\s*```$", "", _raw, flags=re.DOTALL).strip()
        _parsed = json.loads(_clean)
    except Exception:
        _parsed = {"message": _raw, "action": "info", "rerun_pipeline": False, "changes": []}

    _action = _parsed.get("action", "info")
    _msg_text = _parsed.get("message", "")
    _changes = _parsed.get("changes", [])
    _rerun_pipeline = _parsed.get("rerun_pipeline", False)

    _display_msg = _msg_text
    _applied_count = 0

    if _action == "modify" and _changes:
        _cur_ws = st.session_state.get("improved_ws", "")
        _cur_ms = st.session_state.get("improved_ms", "")
        _new_ws, _new_ms, _applied_count = _apply_chat_changes(_changes, _cur_ws, _cur_ms)

        if _applied_count > 0:
            st.session_state["improved_ws"] = _new_ws
            st.session_state["improved_ms"] = _new_ms
            # Queue widget-key updates via pending vars (consumed BEFORE text_area
            # widgets render on the next rerun, avoiding StreamlitAPIException).
            st.session_state["_ws_pending"] = _new_ws
            st.session_state["_ms_pending"] = _new_ms
            # Clear cached export bytes so they regenerate
            for _k in ("fmt_spec", "fmt_docx_bytes", "ms_docx_bytes"):
                st.session_state.pop(_k, None)

            if _rerun_pipeline:
                _display_msg = (
                    f"✅ Applied {_applied_count} change(s). "
                    "Re-running the full enhancement pipeline now — this may take a minute..."
                )
                st.session_state["chat_history"].append({"role": "assistant", "content": _raw, "display": _display_msg})
                st.rerun()
            else:
                _targets = ", ".join(sorted({c.get("target","worksheet") for c in _changes}))
                _display_msg = f"✅ Applied {_applied_count} change(s) to **{_targets}**. {_msg_text}"
        else:
            _display_msg = (
                f"⚠️ Could not find the exact text to replace. {_msg_text}\n\n"
                "Tip: Use the edit boxes above to make the change manually."
            )
    elif _action == "modify" and not _changes:
        _display_msg = _msg_text or "No changes specified."

    st.session_state["chat_history"].append({"role": "assistant", "content": _raw, "display": _display_msg})
    st.rerun()

if st.session_state["chat_history"]:
    if st.button("🗑  Clear chat", key="clear_chat"):
        st.session_state["chat_history"] = []
        st.rerun()
